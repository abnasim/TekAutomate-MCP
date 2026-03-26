#!/usr/bin/env python3
"""
Quick spot check of DPO manual to understand structure
"""

import os
import re
from docx import Document

script_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(script_dir)
docx_path = os.path.join(parent_dir, "MSO-DPO5000-B-DPO7000-C-DPO70000.docx")

if not os.path.exists(docx_path):
    print(f"ERROR: File not found: {docx_path}")
    sys.exit(1)

print(f"Loading: {docx_path}")
doc = Document(docx_path)
print(f"Total paragraphs: {len(doc.paragraphs)}\n")

# Look for first few commands
print("=" * 80)
print("FIRST 50 PARAGRAPHS:")
print("=" * 80)
for i, para in enumerate(doc.paragraphs[:50]):
    text = para.text.strip()
    if not text:
        continue
    
    # Check font info
    font_info = []
    for run in para.runs:
        if run.font and run.font.name:
            font_info.append(run.font.name)
        if run.bold:
            font_info.append("BOLD")
    
    font_str = ", ".join(set(font_info)) if font_info else "default"
    
    # Check if it looks like a command
    looks_like_command = ":" in text and any(c.isupper() for c in text[:20])
    
    marker = ">>> COMMAND?" if looks_like_command else ""
    print(f"[{i:3d}] [{font_str:20s}] {marker} {text[:70]}")

print("\n" + "=" * 80)
print("LOOKING FOR COMMAND PATTERNS:")
print("=" * 80)

# Search for patterns that look like SCPI commands (must have colon)
command_pattern = re.compile(r'^[A-Z][A-Z0-9]*:[A-Z0-9:<>?]+', re.IGNORECASE)
found_commands = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    # Check if it matches command pattern (must have colon and look like SCPI)
    if ':' in text and command_pattern.match(text):
        # Get font info
        fonts = set()
        is_bold = False
        for run in para.runs:
            if run.font and run.font.name:
                fonts.add(run.font.name)
            if run.bold:
                is_bold = True
        
        found_commands.append({
            'para': i,
            'text': text[:100],
            'fonts': list(fonts),
            'bold': is_bold
        })
        
        if len(found_commands) >= 30:  # Show first 30
            break

print(f"\nFound {len(found_commands)} potential commands:")
for cmd in found_commands:
    font_str = ", ".join(cmd['fonts']) if cmd['fonts'] else "default"
    bold_str = " [BOLD]" if cmd['bold'] else ""
    print(f"  Para {cmd['para']:4d}: [{font_str:20s}]{bold_str} {cmd['text']}")

