"""Inspect how examples are structured in Word document"""
from docx import Document
import os
import re

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

# Try both Word files
DOCX_FILENAME_1 = "4-5-6_MSO_Programmer_077189801_RevA (1).docx"
DOCX_FILENAME_2 = "4-5-6_MSO_Programmer_077189801_RevA.docx"

INPUT_DOCX = None
for filename in [DOCX_FILENAME_1, DOCX_FILENAME_2]:
    path = os.path.join(PROJECT_ROOT, filename)
    if os.path.exists(path):
        INPUT_DOCX = path
        break

if not INPUT_DOCX:
    print("Word document not found")
    exit(1)

doc = Document(INPUT_DOCX)

def get_font_name(run):
    if run.font and run.font.name:
        return run.font.name
    return None

def is_courier_new(run):
    font_name = get_font_name(run)
    if font_name:
        font_lower = font_name.lower()
        return "courier" in font_lower and "new" in font_lower
    return False

def is_arial_narrow(run):
    font_name = get_font_name(run)
    if font_name:
        font_lower = font_name.lower()
        return "arial" in font_lower and "narrow" in font_lower
    return False

print("=== Searching for SEARCH:SEARCH1:TRIGger:A:LOGIc:WHEn examples ===\n")

found_examples = False
for para_num, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "SEARCH:SEARCH1:TRIGger:A:LOGIc:WHEn" in text and "FALSE" in text:
        found_examples = True
        print(f"Found at paragraph {para_num}")
        print(f"Full text: {text}")
        print(f"\nFont analysis:")
        
        for i, run in enumerate(para.runs):
            font = get_font_name(run)
            is_courier = is_courier_new(run)
            is_arial = is_arial_narrow(run)
            print(f"  Run {i}: '{run.text[:50]}...' | Font: {font} | Courier: {is_courier} | Arial Narrow: {is_arial}")
        
        # Show the "might return" example
        print(f"\n=== 'might return' example (paragraph {para_num + 1}) ===")
        if para_num + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[para_num + 1]
            next_text = next_para.text.strip()
            print(f"Full text: {next_text}")
            print(f"Font analysis:")
            for i, run in enumerate(next_para.runs):
                font = get_font_name(run)
                is_courier = is_courier_new(run)
                is_arial = is_arial_narrow(run)
                print(f"  Run {i}: '{run.text[:60]}...' | Font: {font} | Courier: {is_courier} | Arial Narrow: {is_arial}")
        
        # Show next paragraph (TRUE indicating...)
        print(f"\n=== Next paragraph (TRUE indicating...) ===")
        if para_num + 2 < len(doc.paragraphs):
            next2_para = doc.paragraphs[para_num + 2]
            next2_text = next2_para.text.strip()
            print(f"Full text: {next2_text}")
            print(f"Font analysis:")
            for i, run in enumerate(next2_para.runs):
                font = get_font_name(run)
                is_courier = is_courier_new(run)
                is_arial = is_arial_narrow(run)
                print(f"  Run {i}: '{run.text[:60]}...' | Font: {font} | Courier: {is_courier} | Arial Narrow: {is_arial}")
        break

if not found_examples:
    print("Example not found")

