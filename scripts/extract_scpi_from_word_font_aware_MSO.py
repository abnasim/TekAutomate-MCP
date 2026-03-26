"""
SCPI Command Extraction Script - MSO 4/5/6 Series
Adapted from DPO extraction script with fixes for:
- Syntax lines with {OPTIONS} on separate lines
- Proper parameter detection
- Conditions capture
"""

import docx
import re
import json
import os
import sys

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

try:
    from command_groups_mapping import COMMAND_TO_GROUP, COMMAND_GROUPS
    HAS_COMMAND_MAPPING = True
    print("Loaded MSO 4/5/6 command group mapping.")
except ImportError:
    print("WARNING: 'command_groups_mapping.py' not found. Using pattern-based command detection.")
    COMMAND_TO_GROUP = {}
    COMMAND_GROUPS = {}
    HAS_COMMAND_MAPPING = False

PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
POSSIBLE_FILES = [
    os.path.join(PROJECT_ROOT, "4-5-6_MSO_Programmer_077189801_RevA (1).docx"),
]
OUTPUT_FILENAME = os.path.join(PROJECT_ROOT, "public", "commands", "mso_2_4_5_6_7.json")

print("Building Master Command Lookup Table...")
# For DPO manual, we'll use pattern-based detection instead of pre-mapped commands
# Commands are detected by SCPI pattern: uppercase letters, colons, optional <x>, optional ?
SCPI_PATTERN = re.compile(r'^[A-Z][A-Z0-9]*:[A-Z0-9:<>?]+(?:\s|$)', re.IGNORECASE)

def is_scpi_command(text):
    """Check if text looks like a SCPI command"""
    if not text or len(text) < 3:
        return False
    # Must have colon and look like SCPI
    if ':' not in text:
        return False
    # Check pattern
    match = SCPI_PATTERN.match(text.strip())
    return match is not None

def get_command_group_from_command(cmd):
    """Get group from mapping file, or infer from command prefix"""
    # First try exact match in mapping
    if HAS_COMMAND_MAPPING:
        # Try with and without query mark
        base_cmd = cmd.replace('?', '').strip()
        if base_cmd in COMMAND_TO_GROUP:
            return COMMAND_TO_GROUP[base_cmd]
        if cmd in COMMAND_TO_GROUP:
            return COMMAND_TO_GROUP[cmd]
        
        # Try parent command matching (e.g., AUXIn:PRObe:DEGAUSS:STATE? -> AUXIn:PRObe:DEGAUSS)
        parts = base_cmd.split(':')
        for i in range(len(parts) - 1, 0, -1):
            parent = ':'.join(parts[:i])
            if parent in COMMAND_TO_GROUP:
                return COMMAND_TO_GROUP[parent]
    
    # Fallback: infer from command prefix
    prefix = cmd.split(':')[0].upper()
    
    # Special handling for MARK commands (should be Search and Mark)
    if prefix == 'MARK':
        return 'Search and Mark'
    
    # Special handling for SELect commands
    if prefix == 'SEL' or prefix == 'SELECT':
        # SELect commands can be in different groups, but most are Search and Mark or Digital
        if 'DIG' in cmd.upper() or 'D<' in cmd:
            return 'Digital'
        return 'Search and Mark'
    
    group_map = {
        'ACQ': 'Acquisition', 'ACQUIRE': 'Acquisition',
        'TRIG': 'Trigger', 'TRIGGER': 'Trigger',
        'CH': 'Vertical', 'CHANNEL': 'Vertical',  # Channels are part of Vertical group
        'HOR': 'Horizontal', 'HORIZONTAL': 'Horizontal',
        'DIS': 'Display control', 'DISPLAY': 'Display control',
        'MEAS': 'Measurement', 'MEASUREMENT': 'Measurement',
        'MATH': 'Math',
        'CURS': 'Cursor', 'CURSOR': 'Cursor',
        'BUS': 'Bus',
        'SAV': 'Save and Recall', 'SAVE': 'Save and Recall',
        'REC': 'Save and Recall', 'RECALL': 'Save and Recall',
        'WAV': 'Waveform Transfer', 'WAVEFORM': 'Waveform Transfer',
        'DAT': 'Waveform Transfer', 'DATA': 'Waveform Transfer',
        'SYST': 'Miscellaneous', 'SYSTEM': 'Miscellaneous',
        'CAL': 'Calibration',
        'DIA': 'Diagnostics', 'DIAG': 'Diagnostics', 'DIAGNOSTICS': 'Diagnostics',
        'ERR': 'Error Detector', 'ERROR': 'Error Detector', 'ERRORDETECTOR': 'Error Detector',
        'EMA': 'E-mail', 'EMAIL': 'E-mail',
        'APP': 'Miscellaneous', 'APPLICATION': 'Miscellaneous',
        'AUX': 'Vertical', 'AUXIN': 'Vertical', 'AUXOUT': 'Miscellaneous',
        'HIS': 'Histogram', 'HISTOGRAM': 'Histogram',
        'LIM': 'Limit Test', 'LIMIT': 'Limit Test',
        'MAS': 'Mask', 'MASK': 'Mask',
        'SEA': 'Search and Mark', 'SEARCH': 'Search and Mark',
        'ZOO': 'Zoom', 'ZOOM': 'Zoom',
        'FIL': 'File system', 'FILE': 'File system',
        'HAR': 'Hard copy', 'HARD': 'Hard copy',
        'LOW': 'Low Speed Serial Trigger', 'LOWS': 'Low Speed Serial Trigger',
        'SAV': 'Save On', 'SAVEON': 'Save On',
        'ROS': 'Miscellaneous', 'ROSC': 'Miscellaneous',  # Reference oscillator
        'IDN': 'Miscellaneous',  # Identification
        'USB': 'Miscellaneous', 'USBTMC': 'Miscellaneous',
        'FPA': 'Miscellaneous', 'FPANEL': 'Miscellaneous',  # Front panel
        'SET': 'Miscellaneous', 'SETUP': 'Miscellaneous',
        'TES': 'Diagnostics', 'TEST': 'Diagnostics',
        'VIS': 'Miscellaneous', 'VISUAL': 'Miscellaneous',
    }
    
    # Check if prefix matches
    if prefix in group_map:
        return group_map[prefix]
    
    # Try partial matches for longer prefixes
    for key, value in group_map.items():
        if prefix.startswith(key):
            return value
    
    return 'Miscellaneous'

SECTION_MAP = {
    "GROUP": "group", "SYNTAX": "syntax", "ARGUMENTS": "arguments",
    "EXAMPLES": "examples", "RELATED COMMANDS": "related", "RELATED": "related",
    "RETURNS": "returns", "CONDITIONS": "conditions"
}


def expand_wildcards(option_list):
    expanded = []
    for opt in option_list:
        opt = opt.strip()
        if not opt:
            continue
        if re.search(r'CH<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"CH{i}" for i in range(1, 9)])
        elif re.search(r'MATH<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"MATH{i}" for i in range(1, 5)])
        elif re.search(r'REF<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"REF{i}" for i in range(1, 5)])
        elif re.search(r'BUS<[xn]>|B<[xn]>', opt, re.IGNORECASE):
            prefix = "B" if "B<" in opt else "BUS"
            expanded.extend([f"{prefix}{i}" for i in range(1, 9)])
        elif re.search(r'MEAS<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"MEAS{i}" for i in range(1, 9)])
        elif re.search(r'SEARCH<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"SEARCH{i}" for i in range(1, 9)])
        elif re.search(r'PLOT<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"PLOT{i}" for i in range(1, 5)])
        elif re.search(r'<NR[123]>|<QString>', opt, re.IGNORECASE):
            continue
        else:
            expanded.append(opt)
    seen = set()
    return [x for x in expanded if not (x in seen or seen.add(x))][:50]


def detect_params(command_header, syntax_list, arguments_text, examples=None):
    params = []
    
    # Extract default value from first SET example (not query)
    example_default = None
    if examples:
        for ex in examples:
            ex_text = ex.get('scpi', '') if isinstance(ex, dict) else str(ex)
            if '?' in ex_text:
                continue  # Skip queries
            
            # Example format: "COMMAND ARG description_starting_with_verb"
            # Find where description starts (verbs like sets, indicates, etc.)
            desc_match = re.search(r'\s+(sets|indicates|specifies|returns|turns|enables|queries|is|might|this)\s+', 
                                   ex_text, re.IGNORECASE)
            if desc_match:
                ex_text = ex_text[:desc_match.start()]  # Keep only SCPI part
            
            # Now extract argument (last token after space)
            if ' ' in ex_text:
                parts = ex_text.rsplit(' ', 1)
                if len(parts) == 2:
                    arg = parts[1].strip()
                    # Validate it looks like a value (number, ON/OFF, quoted string, CH1, etc.)
                    if re.match(r'^-?[\d.]+[eE]?[-+]?\d*$', arg) or \
                       re.match(r'^(ON|OFF|0|1)$', arg, re.IGNORECASE) or \
                       re.match(r'^".*"$', arg) or \
                       re.match(r'^[A-Z]+\d*$', arg, re.IGNORECASE):
                        example_default = arg
                        break
    
    # Pattern 0: Detect mnemonic placeholders in header (e.g., POWer<x>, CH<x>, MEAS<x>)
    # These are editable parts of the command path itself
    header_placeholders = re.findall(r'([A-Z]+)<([xn])>', command_header, re.IGNORECASE)
    for prefix, placeholder in header_placeholders:
        prefix_upper = prefix.upper()
        # Determine the param name and range based on prefix
        if prefix_upper in ['POWER', 'POWER']:
            params.append({"name": "power", "type": "integer", "required": True, 
                          "default": 1, "min": 1, "max": 8, "description": "Power measurement number (1-8)"})
        elif prefix_upper == 'CH':
            params.append({"name": "channel", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Channel number (1-8)"})
        elif prefix_upper == 'MATH':
            params.append({"name": "math", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Math waveform number (1-4)"})
        elif prefix_upper == 'REF':
            params.append({"name": "ref", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Reference waveform number (1-4)"})
        elif prefix_upper == 'MEAS':
            params.append({"name": "meas", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Measurement number (1-8)"})
        elif prefix_upper == 'SEARCH':
            params.append({"name": "search", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Search number (1-8)"})
        elif prefix_upper in ['B', 'BUS']:
            params.append({"name": "bus", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Bus number (1-8)"})
        elif prefix_upper == 'PLOT':
            params.append({"name": "plot", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Plot number (1-4)"})
        elif prefix_upper == 'HISTOGRAM':
            params.append({"name": "histogram", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Histogram number (1-4)"})
        elif prefix_upper == 'CURSOR':
            params.append({"name": "cursor", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 2, "description": "Cursor number (1-2)"})
        elif prefix_upper == 'CALLOUT':
            params.append({"name": "callout", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Callout number (1-8)"})
        elif prefix_upper == 'MASK':
            params.append({"name": "mask", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Mask number (1-4)"})
    
    if syntax_list:
        full_syntax = " ".join(syntax_list)
        
        # Pattern 1: Look for trailing parameters NOT in braces (e.g., "COMMAND CH<x>")
        # Check each syntax entry individually for trailing source params
        for syn in syntax_list:
            if '?' in syn:
                continue  # Skip query syntax
            trailing_match = re.search(r'\s+(CH<[xn]>|MATH<[xn]>|REF<[xn]>)\s*$', syn, re.IGNORECASE)
            if trailing_match:
                trailing_param = trailing_match.group(1).upper()
                expanded = expand_wildcards([trailing_param])
                if expanded:
                    params.append({
                        "name": "source", "type": "enumeration", "required": True,
                        "options": expanded, "default": expanded[0],
                        "description": f"Options: {', '.join(expanded[:5])}..."
                    })
                break  # Only add once
        
        # Pattern 2: Enum in braces {ON|OFF|...}
        enum_matches = re.findall(r'\{([^}]+)\}', full_syntax)
        for match in enum_matches:
            raw_options = [o.strip() for o in match.split('|') if o.strip()]
            # Check if ONLY numeric placeholders (skip - handle later)
            if all(re.match(r'<NR[123]>|<QString>|<Block>', o, re.IGNORECASE) for o in raw_options):
                continue
            # Check if mixed: {ON|OFF|<NR1>} - means ONE param with multiple valid formats
            has_nr1 = any(re.match(r'<NR1>', o, re.IGNORECASE) for o in raw_options)
            # Filter out placeholders, keep only enum values
            enum_only = [o for o in raw_options if not re.match(r'<NR[123]>|<QString>|<Block>', o, re.IGNORECASE)]
            expanded_options = expand_wildcards(enum_only)
            if not expanded_options:
                continue
            param_name = "value"
            if "ON" in expanded_options and "OFF" in expanded_options:
                param_name = "state"
            elif any("CH" in o for o in expanded_options[:5]):
                param_name = "source"
            desc_preview = ", ".join(expanded_options[:5]) + ("..." if len(expanded_options) > 5 else "")
            if has_nr1:
                desc_preview += " (or 0/1)"
            params.append({
                "name": param_name, "type": "enumeration", "required": True,
                "options": expanded_options, "default": expanded_options[0] if expanded_options else None,
                "description": f"Options: {desc_preview}"
            })
            has_enum_arg = True
        
        # Add numeric/string params if no ENUM argument was found
        # (mnemonic params like bus/power don't count - they're for the path, not arguments)
        has_enum_arg = any(p.get('type') == 'enumeration' and p.get('name') not in ['bus', 'power', 'channel', 'math', 'ref', 'meas', 'search', 'plot', 'cursor', 'mask', 'histogram', 'callout', 'source'] for p in params)
        if not has_enum_arg:
            if re.search(r'<NR1>', full_syntax, re.IGNORECASE):
                # Try to get integer default from example
                int_default = None
                if example_default:
                    try:
                        int_default = int(float(example_default))
                    except:
                        pass
                params.append({"name": "value", "type": "integer", "required": True, 
                              "default": int_default, "description": "Integer value"})
            elif re.search(r'<NR[23]>', full_syntax, re.IGNORECASE):
                # Try to get float default from example
                float_default = None
                if example_default:
                    try:
                        float_default = float(example_default)
                    except:
                        pass
                params.append({"name": "value", "type": "float", "required": True,
                              "default": float_default, "description": "Floating point value"})
            elif re.search(r'<QString>', full_syntax, re.IGNORECASE):
                # Use example string as default (strip quotes)
                str_default = None
                if example_default:
                    str_default = example_default.strip('"\'')
                params.append({"name": "label", "type": "string", "required": True,
                              "default": str_default, "description": "Quoted string value"})
    if not params and arguments_text:
        lower_args = arguments_text.lower()
        if "integer" in lower_args:
            params.append({"name": "value", "type": "integer", "required": True})
        elif "float" in lower_args or "nr2" in lower_args or "nr3" in lower_args:
            params.append({"name": "value", "type": "float", "required": True})
        elif "string" in lower_args or "quoted" in lower_args:
            params.append({"name": "label", "type": "string", "required": True})
    return params


def detect_command_type(syntax_list, command_header):
    """Extract set and query syntax, handling combined syntax lines"""
    if not syntax_list:
        if command_header.endswith('?'):
            return {"commandType": "query", "hasQuery": True, "hasSet": False, "querySyntax": command_header, "setSyntax": None}
        else:
            return {"commandType": "both", "hasQuery": True, "hasSet": True, "querySyntax": command_header + "?", "setSyntax": command_header}
    
    set_syntax = ""
    query_syntax = ""
    
    # Use same logic as extract_fast.py to split combined syntax
    for s in syntax_list:
        s = s.strip()
        if not s:
            continue
        
        if "?" in s:
            # Check if line contains BOTH set and query syntax (e.g., "CMD {args} CMD?")
            # First, try splitting on the command header that appears twice
            header = command_header.split(" ")[0].split("?")[0]
            if header and s.count(header) >= 2:
                # Find the second occurrence which is the query
                first_pos = s.find(header)
                second_pos = s.find(header, first_pos + len(header))
                if second_pos > 0:
                    potential_set = s[:second_pos].strip()
                    potential_query = s[second_pos:].strip()
                    if potential_set and not set_syntax:
                        set_syntax = potential_set
                    if potential_query and not query_syntax:
                        query_syntax = potential_query
                    continue
            
            # Try regex pattern
            query_match = re.search(r'\s+([A-Za-z:]+(?:<x>)?[A-Za-z:]*\?)\s*$', s)
            if query_match:
                potential_set = s[:query_match.start()].strip()
                potential_query = query_match.group(1).strip()
                if potential_set and ('{' in potential_set or '<NR' in potential_set or '<QString' in potential_set):
                    if not set_syntax:
                        set_syntax = potential_set
                    if not query_syntax:
                        query_syntax = potential_query
                    continue
            
            # Simple query line
            if not query_syntax:
                query_syntax = s
        else:
            # SET syntax line
            if not set_syntax:
                set_syntax = s
    
    # Fallback defaults
    if not set_syntax and not command_header.endswith("?"):
        set_syntax = command_header.replace("?", "")
    if not query_syntax:
        query_syntax = command_header if command_header.endswith("?") else command_header + "?"
    
    has_query = bool(query_syntax)
    has_set = bool(set_syntax)
    
    if has_query and has_set:
        cmd_type = "both"
    elif has_query:
        cmd_type = "query"
    else:
        cmd_type = "set"
    
    return {"commandType": cmd_type, "hasQuery": has_query, "hasSet": has_set, "querySyntax": query_syntax, "setSyntax": set_syntax}


def find_document():
    for full_path in POSSIBLE_FILES:
        if os.path.exists(full_path):
            return full_path
        return None
    

def clean_text(lines):
    if not lines:
        return None
    text = " ".join(lines).strip()
    return re.sub(r'\s+', ' ', text)


def is_bold(para):
    if para.style and para.style.font and para.style.font.bold:
        return True
    for run in para.runs:
        if run.bold:
            return True
    return False

def is_italic(para):
    if para.style and para.style.font and para.style.font.italic:
        return True
    for run in para.runs:
        if run.italic:
            return True
    return False

def get_font_name(run):
    """Get font name from run"""
    if run.font and run.font.name:
        return run.font.name
    if run.style and hasattr(run.style, 'font') and run.style.font and run.style.font.name:
        return run.style.font.name
    return None

def is_arial_narrow(para):
    """Check if paragraph uses Arial Narrow font"""
    # Check paragraph style first
    if para.style and para.style.font and para.style.font.name:
        if "arial narrow" in para.style.font.name.lower():
            return True
    
    # Check runs
    for run in para.runs:
        font_name = get_font_name(run)
        if font_name and "arial narrow" in font_name.lower():
            return True
    return False

def is_bold(para):
    """Check if paragraph is bold"""
    # Check paragraph style
    if para.style and para.style.font and para.style.font.bold:
        return True
    # Check runs
    for run in para.runs:
        if run.bold:
            return True
    return False

def is_italic(para):
    """Check if paragraph is italic"""
    if para.style and para.style.font and para.style.font.italic:
        return True
    for run in para.runs:
        if run.italic:
            return True
    return False

def is_section_header(para, text):
    """Check if paragraph is a section header (supports both DPO and MSO formats)"""
    # Check if text matches section keywords
    first_word = text.split()[0].upper() if text.split() else ""
    if first_word not in ['GROUP', 'SYNTAX', 'ARGUMENTS', 'EXAMPLES', 'RELATED', 'RETURNS', 'CONDITIONS']:
        return False
    
    # MSO format: Uses "Heading 4" style
    if para.style and para.style.name == 'Heading 4':
        return True
    
    # DPO format: Uses Arial Narrow Bold
    if is_arial_narrow(para) and is_bold(para):
        return True
    
    return False

def is_courier_new(para):
    """Check if paragraph uses Courier New font (for syntax in MSO manuals)"""
    for run in para.runs:
        font_name = get_font_name(run)
        if font_name and "courier new" in font_name.lower():
            return True
    return False

def is_lucida_console(para):
    """Check if paragraph uses Lucida Console font (for syntax in DPO manuals)"""
    for run in para.runs:
        font_name = get_font_name(run)
        if font_name and "lucida console" in font_name.lower():
            return True
    return False

def is_syntax_font(para):
    """Check if paragraph uses syntax font (Courier New for MSO, Lucida Console for DPO)"""
    return is_courier_new(para) or is_lucida_console(para)

def is_times_new_roman(para):
    """Check if paragraph uses Times New Roman font"""
    for run in para.runs:
        font_name = get_font_name(run)
        if font_name and "times new roman" in font_name.lower():
            return True
    return False


def extract_commands(file_path):
    print(f"Loading document: {os.path.basename(file_path)}...")
    doc = docx.Document(file_path)
    commands = []
    current_cmd = None
    current_section = "description"
    buffer = []
    total_paras = len(doc.paragraphs)
    print(f"Scanning {total_paras} paragraphs...")

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this looks like a SCPI command header
        # DPO manual: Command headers are Arial Narrow, Bold
        # Extract first word/command part
        first_part = text.split()[0] if text.split() else ""
        
        # Primary check: Arial Narrow + Bold + SCPI pattern
        is_header_strict = (is_scpi_command(first_part) and 
                           is_arial_narrow(para) and 
                           is_bold(para))
        
        # Fallback: If it's a strong SCPI pattern and bold, accept it
        # (some commands might have font variations)
        is_header_fallback = (is_scpi_command(first_part) and 
                             is_bold(para) and
                             ':' in first_part and
                             len(first_part) > 5)  # Reasonable command length
        
        is_header = is_header_strict or is_header_fallback
        
        if is_header:
            if current_cmd:
                if buffer:
                    current_cmd[current_section].extend(buffer)
                commands.append(current_cmd)
            
            # Extract command (remove any trailing description)
            cmd_header = first_part
            # Try to get full command if it's on one line
            if ' ' in text and is_scpi_command(text.split()[0]):
                # Command might be followed by description
                parts = text.split(None, 1)
                if len(parts) > 0 and is_scpi_command(parts[0]):
                    cmd_header = parts[0]
            
            inferred_group = get_command_group_from_command(cmd_header)
            current_cmd = {
                "header": cmd_header, "mapped_group": inferred_group,
                "description": [], "group": [], "syntax": [], "arguments": [],
                "examples": [], "related": [], "returns": [], "conditions": [], "notes": []
            }
            current_section = "description"
            buffer = []
            continue
        
        if current_cmd:
            first_word_raw = text.split()[0] if text.split() else ""
            first_word = first_word_raw.upper().replace(':', '')
            
            # Check for NOTE section (Arial Narrow Italic)
            if (text.upper().startswith("NOTE") and 
                is_arial_narrow(para) and 
                is_italic(para)):
                current_cmd["notes"].append(text)
                continue
            
            # Check for section headers (Heading 4 style for MSO, Arial Narrow Bold for DPO)
            if text.upper().startswith("RELATED COMMANDS"):
                first_word = "RELATED COMMANDS"
            if first_word in SECTION_MAP and is_section_header(para, text):
                # This is a section header
                if buffer:
                    current_cmd[current_section].extend(buffer)
                current_section = SECTION_MAP[first_word]
                buffer = []
                content_after = text[len(first_word_raw):].strip()
                if content_after.startswith(":"):
                    content_after = content_after[1:].strip()
                # If this is Syntax section and content_after looks like syntax, add to syntax
                if content_after and current_section == "syntax":
                    # For MSO manuals, syntax might have multiple commands on same line
                    # Split by finding SCPI command patterns
                    # Look for patterns like: COMMAND1 COMMAND2? or COMMAND1<param> COMMAND2?
                    syntax_parts = []
                    remaining = content_after
                    
                    # Try to split by finding command headers
                    cmd_header_base = current_cmd["header"].replace('?', '')
                    if cmd_header_base in remaining:
                        # Count occurrences of the command header
                        parts = remaining.split(cmd_header_base)
                        for i, part in enumerate(parts):
                            if i == 0 and not part.strip():
                                continue  # Skip empty first part
                            reconstructed = cmd_header_base + part.strip()
                            if reconstructed and is_scpi_command(reconstructed.split()[0] if reconstructed.split() else ''):
                                syntax_parts.append(reconstructed)
                    
                    # If splitting worked, use parts; otherwise use whole thing
                    if syntax_parts:
                        current_cmd["syntax"].extend(syntax_parts)
                    else:
                        current_cmd["syntax"].append(content_after)
                elif content_after:
                    buffer.append(content_after)
                continue
            
            # Check for syntax lines (Courier New or Lucida Console font)
            if current_section == "syntax" and is_syntax_font(para):
                # This is a syntax line
                current_cmd["syntax"].append(text)
                continue
            
            # ENHANCED: Look for {OPTIONS} patterns in Arguments section or continuation lines
            # Sometimes the options appear in Arguments text or on continuation lines
            if current_section in ["syntax", "arguments"]:
                # Check if this line contains {OPTIONS} pattern
                options_pattern = r'\{([A-Z][A-Za-z0-9]*(?:\|[A-Z][A-Za-z0-9]*)+)\}'
                if re.search(options_pattern, text):
                    # Found options! Check if we need to merge with existing syntax
                    if current_section == "arguments":
                        # Options found in Arguments - try to merge with syntax
                        # Look for the command in syntax and append options
                        if current_cmd["syntax"]:
                            # Check if last syntax line needs options appended
                            last_syntax = current_cmd["syntax"][-1]
                            if '{' not in last_syntax and '|' not in last_syntax:
                                # Syntax line doesn't have options yet - append them
                                options_match = re.search(options_pattern, text)
                                if options_match:
                                    current_cmd["syntax"][-1] = f"{last_syntax} {options_match.group(0)}"
                        else:
                            # No syntax yet, but we found options - create syntax line
                            cmd_header = current_cmd["header"]
                            options_match = re.search(options_pattern, text)
                            if options_match:
                                current_cmd["syntax"].append(f"{cmd_header} {options_match.group(0)}")
                    elif current_section == "syntax":
                        # Options in syntax section but maybe not in Lucida Console
                        # Add it as a syntax line
                        if text not in current_cmd["syntax"]:
                            current_cmd["syntax"].append(text)
                        continue
            
            # ENHANCED: Look for syntax-like patterns even if not in Lucida Console
            # Sometimes syntax appears in regular text (especially continuation lines)
            if current_section == "syntax":
                # Check if text looks like a syntax continuation (contains command pattern or {OPTIONS})
                is_scpi = (is_scpi_command(text.split()[0]) if text.split() else False)
                has_options = re.search(r'\{[A-Z]', text)  # Starts with {OPTIONS}
                
                if is_scpi or has_options:
                    # If it's JUST options on a line by itself, merge with previous syntax
                    if has_options and not is_scpi and text.startswith('{') and current_cmd["syntax"]:
                        # This is an options-only line - merge with previous set syntax
                        last_syntax = current_cmd["syntax"][-1]
                        if '?' not in last_syntax and '{' not in last_syntax:
                            # Append options to the set command
                            current_cmd["syntax"][-1] = f"{last_syntax} {text.strip()}"
                        else:
                            # Can't merge, add as separate line
                            current_cmd["syntax"].append(text)
                    elif text not in current_cmd["syntax"]:
                        current_cmd["syntax"].append(text)
                    continue
            
            # Regular content
            buffer.append(text)
        
        if i % 5000 == 0 and i > 0:
            print(f"  Progress: {i}/{total_paras} ({100*i/total_paras:.1f}%) - {len(commands)} commands", flush=True)

    if current_cmd:
        if buffer:
            current_cmd[current_section].extend(buffer)
        commands.append(current_cmd)
    
    print(f"\nExtracted {len(commands)} commands.")
    
    # Debug: Show first few and last few commands
    if commands:
        print(f"\nFirst 5 commands:")
        for i, cmd in enumerate(commands[:5]):
            print(f"  {i+1}. {cmd['header']}")
        if len(commands) > 5:
            print(f"\nLast 5 commands:")
            for i, cmd in enumerate(commands[-5:], start=len(commands)-4):
                print(f"  {i}. {cmd['header']}")
    
    return commands


def validate_syntax(syntax_list, command_header):
    """Validate and merge syntax lines, combining {OPTIONS} with previous command"""
    base_cmd = command_header.replace('?', '').upper()
    prefix = base_cmd.split(':')[0] if ':' in base_cmd else base_cmd
    valid = []
    
    for i, syn in enumerate(syntax_list):
        syn = syn.strip()
        if not syn:
            continue
            
        # If this is an {OPTIONS} line by itself, try to merge with previous valid line
        if syn.startswith('{') and '|' in syn and valid:
            # Check if previous line is a SET command (not a query)
            if valid and '?' not in valid[-1] and '{' not in valid[-1]:
                # Merge options with previous SET command
                valid[-1] = f"{valid[-1]} {syn}"
            else:
                # Can't merge, add as separate line (will be filtered later if needed)
                valid.append(syn)
        # Keep lines that start with the command prefix
        elif syn.upper().startswith(prefix):
            valid.append(syn)
    
    return valid


def validate_examples(example_lines, command_header):
    prefix = command_header.replace('?', '').upper().split(':')[0]
    valid = []
    for line in example_lines:
        line = line.strip()
        if not line or not line[0].isupper():
            continue
        # Reject description text (common starting words)
        line_lower = line.lower()
        if line_lower.startswith(('this ', 'the ', 'a ', 'an ', 'note', 'see ', 'use ', 'for ', 
                                   'when ', 'if ', 'to ', 'in ', 'on ', 'it ', 'you ', 'requires')):
            continue
        # Must start with the command prefix and contain a colon (SCPI format)
        line_upper = line.upper()
        first_word = line.split()[0] if line.split() else ""
        if first_word.upper().startswith(prefix) and ':' in line:
            valid.append(line)
    return valid


def enhance_syntax_with_arguments(syntax_list, arguments_text, command_header):
    """Enhance syntax by extracting {OPTIONS} from Arguments text if missing from syntax"""
    if not arguments_text or not syntax_list:
        return syntax_list
    
    # Check if syntax already has {OPTIONS}
    full_syntax = " ".join(syntax_list)
    if '{' in full_syntax and '|' in full_syntax:
        return syntax_list  # Already has options
    
    # Look for {OPTIONS} pattern in arguments text
    options_pattern = r'\{([A-Z][A-Za-z0-9]*(?:\|[A-Z][A-Za-z0-9]*)+)\}'
    options_match = re.search(options_pattern, arguments_text)
    
    if options_match:
        # Found options in arguments! Merge into syntax
        options_str = options_match.group(0)
        enhanced_syntax = []
        for syn in syntax_list:
            if '?' in syn:
                enhanced_syntax.append(syn)  # Keep query syntax as-is
            else:
                # Append options to set syntax if not already present
                if '{' not in syn and '|' not in syn:
                    enhanced_syntax.append(f"{syn} {options_str}")
                else:
                    enhanced_syntax.append(syn)
        return enhanced_syntax
    
    # Also look for comma/pipe-separated options in arguments
    # Pattern: "ERRor, DATA, IDANDDATA, EOF, IDentifier, ACKMISS, SOF, FRAMEtype"
    # or "ERRor|DATA|IDANDDATA|EOF|..."
    comma_pipe_pattern = r'\b([A-Z][A-Za-z0-9]+(?:\s*[,|]\s*[A-Z][A-Za-z0-9]+){2,})'
    comma_pipe_match = re.search(comma_pipe_pattern, arguments_text)
    
    if comma_pipe_match:
        options_text = comma_pipe_match.group(1)
        # Convert to {OPTIONS} format
        options_list = re.split(r'[,|]', options_text)
        options_list = [opt.strip() for opt in options_list if opt.strip() and len(opt.strip()) >= 3]
        if len(options_list) >= 2:
            options_str = '{' + '|'.join(options_list) + '}'
            enhanced_syntax = []
            for syn in syntax_list:
                if '?' in syn:
                    enhanced_syntax.append(syn)
                else:
                    if '{' not in syn and '|' not in syn:
                        enhanced_syntax.append(f"{syn} {options_str}")
                    else:
                        enhanced_syntax.append(syn)
            return enhanced_syntax
    
    return syntax_list


def post_process(raw_commands):
    groups_dict = {}
    for cmd in raw_commands:
        header = cmd["header"]
        group_name = cmd["mapped_group"] or "Uncategorized"
        desc = clean_text(cmd["description"]) or ""
        args_text = clean_text(cmd["arguments"])
        validated_syntax = validate_syntax(cmd["syntax"], header)
        
        # ENHANCED: Merge options from Arguments into syntax if missing
        validated_syntax = enhance_syntax_with_arguments(validated_syntax, args_text, header)
        
        validated_examples = validate_examples(cmd["examples"], header)
        cmd_type_info = detect_command_type(validated_syntax, header)
        detected_params = detect_params(header, validated_syntax, args_text, validated_examples)
        
        processed_examples = []
        for line in validated_examples:
            # Pattern: SCPI_COMMAND [ARGUMENT] description_starting_with_verb
            # Description typically starts with: sets, queries, returns, indicates, specifies, turns, enables, might
            desc_start_pattern = r'\s+(sets|queries|returns|indicates|specifies|turns|enables|disables|might|is|will|this)\s+'
            match = re.search(desc_start_pattern, line, re.IGNORECASE)
            if match:
                code = line[:match.start()].strip()
                ex_desc = line[match.start():].strip()
            elif "might return" in line.lower():
                parts = line.split("might return", 1)
                code = parts[0].strip()
                ex_desc = "might return " + parts[1].strip() if len(parts) > 1 else ""
            else:
                # Fallback: split at first lowercase word (but keep uppercase args)
                parts = re.split(r'\s+(?=[a-z])', line, maxsplit=1)
                code = parts[0].strip() if parts else line
                ex_desc = parts[1].strip() if len(parts) > 1 else ""
            processed_examples.append({"scpi": code, "description": ex_desc, "codeExamples": {"scpi": {"code": code}}})
        
        name = header.split(':')[-1].replace('?', '').replace('<x>', '').replace('<n>', '')
        short_desc = desc[:100] + "..." if len(desc) > 100 else desc
        
        manual_entry = {
            "command": header, "header": header.split(':')[0],
            "mnemonics": header.replace('?', '').split(':'),
            "commandType": cmd_type_info["commandType"],
            "hasQuery": cmd_type_info["hasQuery"], "hasSet": cmd_type_info["hasSet"],
            "description": desc, "shortDescription": short_desc, "arguments": args_text,
            "examples": processed_examples,
            "relatedCommands": [c.strip() for c in clean_text(cmd["related"]).split()] if cmd["related"] else [],
            "commandGroup": group_name, "syntaxList": validated_syntax,
            "syntax": {"set": cmd_type_info["setSyntax"], "query": cmd_type_info["querySyntax"]},
            "manualReference": {"section": group_name}, "notes": cmd["notes"]
        }

        final_cmd_obj = {
            "scpi": header, "name": name, "description": desc, "shortDescription": short_desc,
            "group": group_name, "syntax": validated_syntax, "arguments": args_text,
            "params": detected_params, "examples": processed_examples,
            "relatedCommands": manual_entry["relatedCommands"],
            "conditions": clean_text(cmd["conditions"]), "returns": clean_text(cmd["returns"]),
            "notes": cmd["notes"], "example": processed_examples[0]["scpi"] if processed_examples else None,
            "commandType": cmd_type_info["commandType"],
            "hasQuery": cmd_type_info["hasQuery"], "hasSet": cmd_type_info["hasSet"],
            "_manualEntry": manual_entry
        }

        if group_name not in groups_dict:
            # Get description from COMMAND_GROUPS if available
            group_desc = ""
            if HAS_COMMAND_MAPPING and group_name in COMMAND_GROUPS:
                group_desc = COMMAND_GROUPS[group_name].get("description", "")
            groups_dict[group_name] = {
                "name": group_name,
                "description": group_desc,
                "commands": []
            }
        groups_dict[group_name]["commands"].append(final_cmd_obj)

    return {
        "version": "2.0", "manual": "4/5/6 Series MSO Programmer Manual",
        "groups": groups_dict,
        "metadata": {"total_commands": len(raw_commands), "total_groups": len(groups_dict)}
    }


if __name__ == "__main__":
    found_file = find_document()
    if not found_file:
        print("ERROR: Word document not found.")
        sys.exit(1)
    
    print("Extracting (Golden Key Verification Mode)...")
    raw_data = extract_commands(found_file)
    print(f"\nExtracted {len(raw_data)} commands matched against Mapping File.")
    
    final_json = post_process(raw_data)
    os.makedirs(os.path.dirname(OUTPUT_FILENAME), exist_ok=True)
    
    with open(OUTPUT_FILENAME, "w", encoding="utf-8") as f:
        json.dump(final_json, f, indent=2, ensure_ascii=False)
    
    print(f"\nSUCCESS! Output saved to: {OUTPUT_FILENAME}")
    print(f"  Commands: {final_json['metadata']['total_commands']}")
    print(f"  Groups: {final_json['metadata']['total_groups']}")

