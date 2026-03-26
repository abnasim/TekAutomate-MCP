"""Find all occurrences of PLOT:PLOT<x>:RAILNUM in the document"""
import os
import sys

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
DOCX_PATH = os.path.join(PROJECT_ROOT, "4-5-6_MSO_Programmer_077189801_RevA (1).docx")

print("Loading document...")
doc = Document(DOCX_PATH)

print("\nSearching for 'RAILNUM' in paragraphs...")
count = 0
for i, para in enumerate(doc.paragraphs):
    if "RAILNUM" in para.text.upper():
        count += 1
        print(f"[{i}] {para.text[:80]}...")

print(f"\nTotal paragraph occurrences: {count}")

print("\nSearching in tables...")
table_count = 0
for t, table in enumerate(doc.tables):
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            if "RAILNUM" in cell.text.upper():
                table_count += 1
                print(f"Table {t}, Row {r}, Cell {c}: {cell.text[:80]}...")

print(f"\nTotal table occurrences: {table_count}")









