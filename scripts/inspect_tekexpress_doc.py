"""
Quick inspection script to understand TekExpress manual structure
"""
import docx
import re
import os
import sys

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
DOCX_FILE = os.path.join(PROJECT_ROOT, "TekExpress_USB4Tx_UserManual_EN-US_077-1702-04_077170204.docx")

if not os.path.exists(DOCX_FILE):
    print(f"ERROR: File not found: {DOCX_FILE}")
    sys.exit(1)

print(f"Loading: {os.path.basename(DOCX_FILE)}")
doc = docx.Document(DOCX_FILE)

# Extract group name from filename - fix regex for USB4Tx
filename = os.path.basename(DOCX_FILE)
# Try to match TekExpress_<SuiteName>_ pattern
group_match = re.search(r'TekExpress_([A-Z0-9a-z]+)', filename, re.IGNORECASE)
group_name = group_match.group(1) if group_match else "TekExpress"
print(f"\nExtracted group name from filename: {group_name}")

# Find first command and examine structure around it
print("\n=== Examining structure around first TEKEXP command ===")
tekexp_pattern = re.compile(r'TEKEXP:[A-Z]+', re.IGNORECASE)
first_cmd_idx = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if tekexp_pattern.search(text) and '(Set)' in text or '(Query)' in text:
        first_cmd_idx = i
        break

if first_cmd_idx:
    print(f"\nFirst command found at paragraph {first_cmd_idx}")
    print("\n=== Context around first command (20 paragraphs before and after) ===")
    start = max(0, first_cmd_idx - 20)
    end = min(len(doc.paragraphs), first_cmd_idx + 20)
    
    for i in range(start, end):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if not text:
            continue
        
        # Mark the command
        marker = ">>> COMMAND <<<" if i == first_cmd_idx else ""
        
        # Check for section markers
        text_lower = text.lower()
        is_section = False
        if text_lower in ["syntax", "command arguments", "arguments", "returns", "examples"]:
            marker = ">>> SECTION <<<"
            is_section = True
        
        # Check if it's a command
        is_cmd = tekexp_pattern.search(text)
        if is_cmd:
            marker = ">>> COMMAND <<<"
        
        print(f"{i:4d} {marker:20s} {text[:100]}")

# Look for tables near commands
print("\n=== Looking for argument tables (checking tables near commands) ===")
# Find a command that likely has arguments
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'TEKEXP:SELECT TEST' in text.upper():
        print(f"\nFound TEKEXP:SELECT TEST at paragraph {i}")
        # Look for nearby tables
        # Tables are separate from paragraphs, so we need to find them differently
        # Let's check if there are paragraphs mentioning "Command arguments" or "Argument"
        for j in range(max(0, i-10), min(len(doc.paragraphs), i+30)):
            p = doc.paragraphs[j]
            pt = p.text.strip()
            if 'command arguments' in pt.lower() or 'argument' in pt.lower():
                print(f"  Para {j}: {pt[:80]}")
        break

# Check a specific table that might be argument table
print("\n=== Checking tables that might contain arguments ===")
for table_idx, table in enumerate(doc.tables):
    if table_idx < 10:  # Skip first few (document structure)
        continue
    if table_idx > 30:  # Don't check too many
        break
    
    # Check first row for argument-related headers
    if len(table.rows) > 0:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if any('argument' in cell.lower() or 'testname' in cell.lower() or 'value' in cell.lower() 
               for cell in first_row):
            print(f"\nTable {table_idx}: {len(table.rows)} rows, {len(table.columns)} columns")
            print(f"  First row: {' | '.join(first_row)}")
            # Show a few data rows
            for row_idx, row in enumerate(table.rows[1:6]):
                cells = [cell.text.strip() for cell in row.cells]
                print(f"  Row {row_idx+1}: {' | '.join(cells[:3])}")

print("\n=== Inspection complete ===")
