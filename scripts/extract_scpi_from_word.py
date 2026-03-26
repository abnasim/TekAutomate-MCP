"""
Enhanced SCPI Command Extraction Script from Word Document
Word format preserves text structure much better than PDF extraction
"""

import json
import re
import sys
import os

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Install it with: pip install python-docx")
    sys.exit(1)

# Add scripts directory to path to import command groups mapping
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from command_groups_mapping import COMMAND_GROUPS

# ================= CONFIGURATION =================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

DOCX_FILENAME = "4-5-6_MSO_Programmer_077189801_RevA.docx"
INPUT_DOCX = os.path.join(PROJECT_ROOT, DOCX_FILENAME)

if not os.path.exists(INPUT_DOCX):
    print(f"ERROR: Word document not found: {INPUT_DOCX}")
    sys.exit(1)

OUTPUT_JSON = os.path.join(PROJECT_ROOT, "public", "commands", "mso_commands_final.json")

# Regex patterns
CMD_PATTERN = re.compile(
    r'^([*][A-Za-z]{2,}\??)$|'  # Star commands
    r'^([A-Za-z][A-Za-z0-9<>]*:[A-Za-z0-9<>:]+(?:\?)?)$'  # Standard commands
)

def get_group_for_command(cmd):
    """Find which group a command belongs to using the mapping."""
    if not cmd:
        return None
    
    normalized = cmd.upper().replace('?', '').replace('<N>', '<n>').replace('<X>', '<x>')
    
    # Normalize special patterns with <x> in the middle (before "Val" or "Voltage" or "VOLTage")
    normalized = re.sub(r'PG(\d+)VAL', r'PG<x>VAL', normalized)
    normalized = re.sub(r'PW(\d+)VAL', r'PW<x>VAL', normalized)
    normalized = re.sub(r'AMP(\d+)VAL', r'AMP<x>VAL', normalized)
    normalized = re.sub(r'FREQ(\d+)VAL', r'FREQ<x>VAL', normalized)
    normalized = re.sub(r'SPAN(\d+)VAL', r'SPAN<x>VAL', normalized)
    normalized = re.sub(r'RIPPLEFREQ(\d+)VAL', r'RIPPLEFREQ<x>VAL', normalized)
    normalized = re.sub(r'MAXG(\d+)VOLTAGE', r'MAXG<x>VOLTAGE', normalized)
    normalized = re.sub(r'OUTPUT(\d+)VOLTAGE', r'OUTPUT<x>VOLTAGE', normalized)
    
    # Normalize view patterns
    normalized = re.sub(r'WAVEVIEW\d+', r'WAVEVIEW<x>', normalized)
    normalized = re.sub(r'PLOTVIEW\d+', r'PLOTVIEW<x>', normalized)
    normalized = re.sub(r'MATHFFTVIEW\d+', r'MATHFFTVIEW<x>', normalized)
    normalized = re.sub(r'REFFFTVIEW\d+', r'REFFFTVIEW<x>', normalized)
    normalized = re.sub(r'SPECVIEW\d+', r'SPECVIEW<x>', normalized)
    
    for group_name, group_data in COMMAND_GROUPS.items():
        commands_list = group_data.get("commands", [])
        for mapped_cmd in commands_list:
            mapped_normalized = mapped_cmd.upper().replace('?', '').replace('<N>', '<n>').replace('<X>', '<x>')
            
            if normalized == mapped_normalized:
                return group_name
            
            if normalized.startswith(mapped_normalized) or mapped_normalized.startswith(normalized):
                base_normalized = re.sub(r'[<{][nx][>}]', '', normalized)
                base_mapped = re.sub(r'[<{][nx][>}]', '', mapped_normalized)
                if base_normalized == base_mapped:
                    return group_name
    
    return None

def clean_text(text_or_list):
    """Clean extracted text - Word format should be much cleaner."""
    if not text_or_list:
        return None
    
    # Handle both string and list inputs
    if isinstance(text_or_list, list):
        text = "\n".join([line.strip() for line in text_or_list if line.strip()])
    else:
        text = text_or_list
    
    if not text:
        return None
    
    # Word format should be clean, but remove excessive whitespace
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n\s*\n', '\n', text)
    return text.strip() if text.strip() else None

def extract_mso_commands(docx_path):
    print(f"Opening Word document: {docx_path}...")
    print(f"Using {len(COMMAND_GROUPS)} command groups for validation")
    
    doc = Document(docx_path)
    commands = {}
    current_cmd = None
    state = "SEARCHING"
    buffer = []
    
    def save_current_section():
        """Saves buffer content to the appropriate field."""
        nonlocal buffer, state
        if not current_cmd:
            return
        
        text_content = clean_text(buffer)
        if not text_content:
            return
        
        if state == "DESCRIPTION":
            current_cmd["description"] = text_content
        elif state == "CONDITIONS":
            current_cmd["conditions"] = text_content
        elif state == "GROUP":
            extracted_group = text_content.strip()
            mapped_group = get_group_for_command(current_cmd["scpi"])
            current_cmd["group"] = mapped_group or extracted_group
        elif state == "SYNTAX":
            syntax_lines = [line.strip() for line in buffer if line.strip()]
            if syntax_lines:
                if not current_cmd.get("syntax"):
                    current_cmd["syntax"] = []
                current_cmd["syntax"].extend(syntax_lines)
        elif state == "ARGUMENTS":
            current_cmd["arguments"] = text_content
        elif state == "EXAMPLES":
            if not current_cmd.get("examples"):
                current_cmd["examples"] = ""
            if current_cmd["examples"]:
                current_cmd["examples"] += "\n" + text_content
            else:
                current_cmd["examples"] = text_content
        elif state == "RELATED":
            clean = text_content.replace(',', ' ').replace('\n', ' ')
            related = [cmd.strip() for cmd in clean.split() if cmd.strip() and CMD_PATTERN.match(cmd.strip())]
            if related:
                if not current_cmd.get("relatedCommands"):
                    current_cmd["relatedCommands"] = []
                current_cmd["relatedCommands"].extend(related)
        elif state == "RETURNS":
            current_cmd["returns"] = text_content
        
        buffer = []
    
    def finalize_command():
        """Saves the current command and resets."""
        nonlocal current_cmd, state, buffer
        
        if not current_cmd:
            return
        
        save_current_section()
        
        if not current_cmd.get("group"):
            mapped_group = get_group_for_command(current_cmd["scpi"])
            if mapped_group:
                current_cmd["group"] = mapped_group
        
        # Normalize fields
        if not current_cmd.get("relatedCommands"):
            current_cmd["relatedCommands"] = None
        if not current_cmd.get("conditions"):
            current_cmd["conditions"] = None
        if not current_cmd.get("returns"):
            current_cmd["returns"] = None
        if not current_cmd.get("syntax"):
            current_cmd["syntax"] = []
        if not current_cmd.get("examples"):
            current_cmd["examples"] = None
        
        # Deduplication
        cmd_key = current_cmd["scpi"]
        if cmd_key in commands:
            existing = commands[cmd_key]
            for key in ["description", "group", "arguments", "examples", "conditions", "returns"]:
                if not existing.get(key) and current_cmd.get(key):
                    existing[key] = current_cmd[key]
            if current_cmd.get("syntax"):
                existing["syntax"].extend([s for s in current_cmd["syntax"] if s not in existing["syntax"]])
            if current_cmd.get("relatedCommands"):
                if not existing.get("relatedCommands"):
                    existing["relatedCommands"] = []
                existing["relatedCommands"].extend([r for r in current_cmd["relatedCommands"] if r not in existing["relatedCommands"]])
        else:
            commands[cmd_key] = current_cmd
        
        current_cmd = None
        state = "SEARCHING"
        buffer = []
    
    # Extract text from Word document
    # Word documents have paragraphs AND tables - need to check both
    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    print(f"Processing {total_paragraphs} paragraphs and {total_tables} tables...")
    
    # First, process all paragraphs
    for para_num, paragraph in enumerate(doc.paragraphs):
        line = paragraph.text.strip()
        if not line:
            continue
        
        line_lower = line.lower()
        
        # State switching keywords
        if line_lower.startswith("conditions"):
            save_current_section()
            state = "CONDITIONS"
            buffer = []
            continue
        elif line_lower.startswith("group"):
            save_current_section()
            state = "GROUP"
            buffer = []
            continue
        elif line_lower.startswith("syntax"):
            save_current_section()
            state = "SYNTAX"
            buffer = []
            continue
        elif line_lower.startswith("related commands") or (line_lower.startswith("related") and "command" in line_lower):
            save_current_section()
            state = "RELATED"
            buffer = []
            continue
        elif line_lower.startswith("arguments") or line_lower.startswith("args"):
            save_current_section()
            state = "ARGUMENTS"
            buffer = []
            continue
        elif line_lower.startswith("examples") or line_lower.startswith("example"):
            save_current_section()
            state = "EXAMPLES"
            buffer = []
            continue
        elif line_lower.startswith("returns"):
            save_current_section()
            state = "RETURNS"
            buffer = []
            continue
        
        # Check for command header
        # Be more lenient - allow detection after sections complete
        words = line.split()
        if words:
            first_word = words[0]
            # Allow command detection in more cases:
            # - When searching (no current command)
            # - When in DESCRIPTION but no description captured yet
            # - When in other states but buffer is empty (section transition)
            can_detect = (state == "SEARCHING" or 
                         (state == "DESCRIPTION" and current_cmd and not current_cmd.get("description")) or
                         (state not in ["SYNTAX", "EXAMPLES"] and (not buffer or len(buffer) == 0)))
            
            if CMD_PATTERN.match(first_word) and can_detect:
                finalize_command()
                
                # Start new command
                current_cmd = {
                    "scpi": first_word,
                    "description": None,
                    "conditions": None,
                    "group": None,
                    "syntax": [],
                    "relatedCommands": None,
                    "arguments": None,
                    "examples": None,
                    "returns": None
                }
                
                # Check if description is on same line
                if len(words) > 1:
                    desc_text = " ".join(words[1:])
                    # Check if it's a real description (not another command or syntax)
                    if desc_text and not CMD_PATTERN.match(desc_text.split()[0] if desc_text.split() else "") and len(desc_text) < 200:
                        if not (desc_text.count(':') > 3 or desc_text.count('{') > 0 or desc_text.count('|') > 0):
                            # This looks like a description
                            current_cmd["description"] = clean_text([desc_text])
                            state = "SEARCHING"
                        else:
                            # Likely syntax, wait for description on next line
                            state = "DESCRIPTION"
                    else:
                        # No description on same line, expect it on next line
                        state = "DESCRIPTION"
                else:
                    # Command only, description should be on next line
                    state = "DESCRIPTION"
                
                buffer = []
                continue
        
        # If we're in DESCRIPTION state and have a command, this line is likely the description
        if state == "DESCRIPTION" and current_cmd and not current_cmd.get("description"):
            # Check if this line looks like a description (not a command, not a section header)
            if not any(line_lower.startswith(kw) for kw in ["syntax", "arguments", "examples", "group", "related", "returns", "conditions"]):
                if not CMD_PATTERN.match(line.split()[0] if line.split() else ""):
                    # This is likely the description
                    current_cmd["description"] = clean_text([line])
                    state = "SEARCHING"
                    buffer = []
                    continue
        
        # Accumulate content
        if state != "SEARCHING" and current_cmd:
            buffer.append(line)
        
        if para_num % 500 == 0:
            print(f"  Processed paragraph {para_num + 1}/{total_paragraphs}...")
    
    # Process tables - commands might be in table format
    if total_tables > 0:
        print(f"\nProcessing {total_tables} tables...")
        for table_num, table in enumerate(doc.tables):
            # Process each row in the table
            for row in table.rows:
                # Check if this row has command structure (usually 2+ columns)
                if len(row.cells) >= 2:
                    cmd_cell = row.cells[0].text.strip()
                    desc_cell = row.cells[1].text.strip() if len(row.cells) > 1 else ""
                    
                    # Check if first cell is a command
                    words = cmd_cell.split()
                    if words:
                        first_word = words[0]
                        if CMD_PATTERN.match(first_word):
                            finalize_command()
                            
                            current_cmd = {
                                "scpi": first_word,
                                "description": desc_cell if desc_cell and not CMD_PATTERN.match(desc_cell) else None,
                                "conditions": None,
                                "group": None,
                                "syntax": [],
                                "relatedCommands": None,
                                "arguments": None,
                                "examples": None,
                                "returns": None
                            }
                            
                            if desc_cell and not CMD_PATTERN.match(desc_cell):
                                state = "SEARCHING"
                            else:
                                state = "DESCRIPTION"
                            
                            buffer = []
                            continue
                
                # Also process all cell text for section content
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue
                    
                    # Process cell text line by line
                    for line in cell_text.split('\n'):
                        line = line.strip()
                        if not line:
                            continue
                        
                        line_lower = line.lower()
                        
                        # State switching keywords
                        if line_lower.startswith("conditions"):
                            save_current_section()
                            state = "CONDITIONS"
                            buffer = []
                            continue
                        elif line_lower.startswith("group"):
                            save_current_section()
                            state = "GROUP"
                            buffer = []
                            continue
                        elif line_lower.startswith("syntax"):
                            save_current_section()
                            state = "SYNTAX"
                            buffer = []
                            continue
                        elif line_lower.startswith("related commands") or (line_lower.startswith("related") and "command" in line_lower):
                            save_current_section()
                            state = "RELATED"
                            buffer = []
                            continue
                        elif line_lower.startswith("arguments") or line_lower.startswith("args"):
                            save_current_section()
                            state = "ARGUMENTS"
                            buffer = []
                            continue
                        elif line_lower.startswith("examples") or line_lower.startswith("example"):
                            save_current_section()
                            state = "EXAMPLES"
                            buffer = []
                            continue
                        elif line_lower.startswith("returns"):
                            save_current_section()
                            state = "RETURNS"
                            buffer = []
                            continue
                        
                        # Check for command header in table cells
                        words = line.split()
                        if words:
                            first_word = words[0]
                            if CMD_PATTERN.match(first_word) and state != "SYNTAX" and state != "EXAMPLES":
                                finalize_command()
                                
                                current_cmd = {
                                    "scpi": first_word,
                                    "description": None,
                                    "conditions": None,
                                    "group": None,
                                    "syntax": [],
                                    "relatedCommands": None,
                                    "arguments": None,
                                    "examples": None,
                                    "returns": None
                                }
                                
                                if len(words) > 1:
                                    desc_text = " ".join(words[1:])
                                    if desc_text and not CMD_PATTERN.match(desc_text) and len(desc_text) < 200:
                                        if not (desc_text.count(':') > 3):
                                            current_cmd["description"] = clean_text([desc_text])
                                            state = "SEARCHING"
                                        else:
                                            state = "DESCRIPTION"
                                    else:
                                        state = "DESCRIPTION"
                                else:
                                    state = "DESCRIPTION"
                                
                                buffer = []
                                continue
                        
                        # Accumulate content
                        if state != "SEARCHING" and current_cmd:
                            buffer.append(line)
            
            if table_num % 10 == 0 and total_tables > 10:
                print(f"  Processed table {table_num + 1}/{total_tables}...")
    
    finalize_command()
    
    commands_list = list(commands.values())
    
    # Post-process: assign groups
    unassigned = 0
    for cmd in commands_list:
        if not cmd.get("group"):
            mapped_group = get_group_for_command(cmd["scpi"])
            if mapped_group:
                cmd["group"] = mapped_group
            else:
                unassigned += 1
    
    print(f"\nExtraction complete!")
    print(f"  Total commands extracted: {len(commands_list)}")
    print(f"  Commands with assigned groups: {len(commands_list) - unassigned}")
    print(f"  Commands without groups: {unassigned}")
    
    return commands_list

# Copy all the post-processing functions from extract_scpi_enhanced.py
def detect_params(scpi_command):
    """Detect editable parameters in SCPI command like {n}, <n>, <x>"""
    params = []
    
    pattern1 = re.compile(r'\{(\w+)\}')
    matches1 = pattern1.findall(scpi_command)
    for param_name in matches1:
        params.append({
            "name": param_name,
            "type": "number" if param_name in ['n', 'x', 'y', 'z'] else "text",
            "default": 1 if param_name in ['n', 'x'] else None,
            "required": True
        })
    
    pattern2 = re.compile(r'<(\w+)>')
    matches2 = pattern2.findall(scpi_command)
    for param_name in matches2:
        if not any(p["name"] == param_name for p in params):
            params.append({
                "name": param_name,
                "type": "number" if param_name in ['n', 'x', 'y', 'z'] else "text",
                "default": 1 if param_name in ['n', 'x'] else None,
                "required": True
            })
    
    pattern3 = re.compile(r'(CH|REF|MATH|MEAS|BUS|B|CURSOR|ZOOM|SEARCH|PLOT|WAVEView|PLOTView)(<(\w+)>|\{(\w+)\})', re.IGNORECASE)
    matches3 = pattern3.findall(scpi_command)
    for match in matches3:
        param_name = match[2] or match[3]
        if param_name and not any(p["name"] == param_name for p in params):
            params.append({
                "name": param_name,
                "type": "number",
                "default": 1,
                "required": True
            })
    
    return params

def generate_name(scpi, description, short_description):
    """Generate a user-friendly name for the command"""
    if description and len(description) < 100:
        desc = description.replace('Sets or queries ', '').replace('sets or queries ', '')
        if desc:
            return desc[0].upper() + desc[1:] if desc else scpi
    
    if short_description and len(short_description) < 100:
        return short_description[0].upper() + short_description[1:] if short_description else scpi
    
    scpi_parts = scpi.split(':')
    if scpi_parts:
        last_part = scpi_parts[-1].replace('?', '').replace('<n>', 'n').replace('<x>', 'x').replace('{n}', 'n').replace('{x}', 'x')
        words = re.findall(r'[A-Z][a-z]*|[A-Z]+(?=[A-Z]|$)', last_part)
        if words:
            return ' '.join(words)
    
    return scpi

# ================= EXECUTION =================
if __name__ == "__main__":
    data = extract_mso_commands(INPUT_DOCX)
    
    # Organize commands by groups (same as extract_scpi_enhanced.py)
    groups_dict = {}
    ungrouped = []
    
    for cmd in data:
        # Ensure arguments is either null or an empty array
        if cmd.get("arguments") is not None and not isinstance(cmd.get("arguments"), list):
            cmd["arguments"] = None
        
        # Clean and fix description
        description = cmd.get("description")
        if description:
            first_sentence = description.split('.')[0].strip()
            if first_sentence:
                cmd["shortDescription"] = first_sentence[:80] if len(first_sentence) > 80 else first_sentence
            else:
                cmd["shortDescription"] = description[:80] if len(description) > 80 else description
        else:
            scpi_parts = cmd.get("scpi", "").split(':')
            if scpi_parts:
                last_part = scpi_parts[-1].replace('?', '').replace('<n>', 'n').replace('<x>', 'x')
                words = re.findall(r'[A-Z][a-z]*|[A-Z]+(?=[A-Z]|$)', last_part)
                readable = ' '.join(words).lower() if words else last_part.lower()
                cmd["description"] = f"Sets or queries {readable}"
                cmd["shortDescription"] = f"Sets or queries {readable}"
        
        # Generate name field
        cleaned_description = cmd.get("description", "")
        cleaned_short = cmd.get("shortDescription", "")
        cmd["name"] = generate_name(cmd.get("scpi", ""), cleaned_description, cleaned_short)
        
        # Detect and add params array
        scpi_cmd = cmd.get("scpi", "")
        detected_params = detect_params(scpi_cmd)
        cmd["params"] = detected_params if detected_params else []
        
        # Store examples as string for main command
        examples_str = cmd.get("examples")
        if isinstance(examples_str, str) and examples_str:
            cmd["example"] = examples_str.split('\n')[0] if examples_str else None
        else:
            cmd["example"] = None
        
        # Create examples array for manualEntry
        examples_array = []
        if isinstance(examples_str, str) and examples_str:
            example_lines = [line.strip() for line in examples_str.split('\n') if line.strip()]
            examples_array = [
                {
                    "description": f"Example {i+1}",
                    "codeExamples": {
                        "scpi": {"code": line}
                    }
                }
                for i, line in enumerate(example_lines)
            ]
        
        # Create manualEntry structure
        scpi_cmd = cmd.get("scpi", "")
        cmd["_manualEntry"] = {
            "command": scpi_cmd,
            "header": scpi_cmd.split(' ')[0].split('?')[0] if scpi_cmd else "",
            "mnemonics": scpi_cmd.split(' ')[0].split(':') if scpi_cmd else [],
            "commandType": "query" if scpi_cmd.endswith('?') else ("both" if '?' in scpi_cmd else "set"),
            "description": cleaned_description or cmd.get("shortDescription", ""),
            "shortDescription": cmd.get("shortDescription", ""),
            "arguments": None,
            "examples": examples_array,
            "relatedCommands": cmd.get("relatedCommands") if isinstance(cmd.get("relatedCommands"), list) else [],
            "commandGroup": cmd.get("group", ""),
            "syntax": {
                "set": scpi_cmd.replace('?', '') if scpi_cmd else "",
                "query": scpi_cmd if scpi_cmd and scpi_cmd.endswith('?') else (scpi_cmd + '?' if scpi_cmd else "")
            } if scpi_cmd else None,
            "manualReference": {"section": cmd.get("group", "")}
        }
        
        # Remove examples field from main command (use example instead)
        if "examples" in cmd:
            del cmd["examples"]
        
        group_name = cmd.get("group")
        if group_name:
            if group_name not in groups_dict:
                groups_dict[group_name] = {
                    "description": COMMAND_GROUPS.get(group_name, {}).get("description", ""),
                    "commands": []
                }
            groups_dict[group_name]["commands"].append(cmd)
        else:
            ungrouped.append(cmd)
    
    # Add ungrouped commands to Miscellaneous
    if ungrouped:
        if "Miscellaneous" not in groups_dict:
            groups_dict["Miscellaneous"] = {
                "description": COMMAND_GROUPS.get("Miscellaneous", {}).get("description", "Commands that do not fit into other categories."),
                "commands": []
            }
        groups_dict["Miscellaneous"]["commands"].extend(ungrouped)
    
    # Wrap in final structure
    final_output = {
        "version": "1.0.0",
        "manual": {
            "title": "4-5-6 Series MSO Programmer Manual",
            "file": "4-5-6_MSO_Programmer_077189801_RevA.docx",
            "revision": "A",
            "models": ["MSO4XB", "MSO5XB", "MSO58LP", "MSO6XB", "LPD64"],
            "families": ["MSO4", "MSO5", "MSO6", "MSO7"]
        },
        "groups": groups_dict,
        "metadata": {
            "total_commands": len(data),
            "total_groups": len(groups_dict),
            "command_groups_count": len(COMMAND_GROUPS),
            "extraction_date": "2024",
            "note": "Some commands may not be available on all instrument models. Some commands require specific options to be installed."
        }
    }
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(OUTPUT_JSON), exist_ok=True)
    
    print(f"\nSaving to {OUTPUT_JSON}...")
    print(f"  Organized into {len(groups_dict)} groups")
    with open(OUTPUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(final_output, f, indent=2, ensure_ascii=False)
    
    print("Extraction Complete!")
    print(f"\nThe file is ready to use in TekAutomate!")
    print(f"Location: {OUTPUT_JSON}")

