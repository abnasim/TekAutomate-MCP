"""
Debug extraction for a single command to see what's happening
"""
import os
import sys
import re

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

# Target command to debug
TARGET_CMD = "PLOT:PLOT<x>:RAILNUM"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
DOCX_PATH = os.path.join(PROJECT_ROOT, "4-5-6_MSO_Programmer_077189801_RevA (1).docx")

def get_font_name(run):
    """Get font name from a run, checking run and style."""
    if run.font and run.font.name:
        return run.font.name
    if run.style and hasattr(run.style, 'font') and run.style.font and run.style.font.name:
        return run.style.font.name
    return None

print(f"Loading document: {DOCX_PATH}")
doc = Document(DOCX_PATH)
print(f"Loaded. {len(doc.paragraphs)} paragraphs")

# Find the target command and print 30 paragraphs after it
found = False
capture_count = 0
MAX_CAPTURE = 40

print(f"\nSearching for: {TARGET_CMD}\n")
print("="*80)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if not found:
        # Look for our target command
        if TARGET_CMD.replace("<x>", "") in text.replace("<x>", "").upper() or "RAILNUM" in text.upper():
            if "PLOT" in text.upper() and "RAILNUM" in text.upper():
                found = True
                print(f"\n>>> FOUND TARGET at paragraph {i} <<<\n")
    
    if found:
        capture_count += 1
        
        # Analyze fonts in this paragraph - show ALL font info
        fonts_used = []
        for run in para.runs:
            font = get_font_name(run)
            run_text = run.text.strip()
            if run_text:
                font_info = font if font else "NO_FONT_DETECTED"
                # Check style too
                style_name = run.style.name if run.style else "no_style"
                fonts_used.append(f"{font_info} (style:{style_name}): '{run_text[:50]}'")
        
        print(f"[{i}] TEXT: {text[:100]}{'...' if len(text) > 100 else ''}")
        if fonts_used:
            for f in fonts_used[:3]:
                print(f"      FONT: {f}")
        print()
        
        if capture_count >= MAX_CAPTURE:
            break

if not found:
    print("Target command not found! Let's search more broadly...")
    for i, para in enumerate(doc.paragraphs):
        if "RAILNUM" in para.text.upper():
            print(f"[{i}] {para.text[:100]}")

