"""
SCPI Command Extraction Script - Golden Key Verification Mode
"""

import docx
import re
import json
import os
import sys

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

try:
    from command_groups_mapping import COMMAND_GROUPS
except ImportError:
    print("CRITICAL ERROR: 'command_groups_mapping.py' not found.")
    sys.exit(1)

PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
POSSIBLE_FILES = [
    os.path.join(PROJECT_ROOT, "4-5-6_MSO_Programmer_077189801_RevA (1).docx"),
    os.path.join(PROJECT_ROOT, "4-5-6_MSO_Programmer_077189801_RevA.docx"),
]
OUTPUT_FILENAME = os.path.join(PROJECT_ROOT, "public", "commands", "mso_commands_final.json")

print("Building Master Command Lookup Table...")
MASTER_CMD_MAP = {}

for group_name, data in COMMAND_GROUPS.items():
    for cmd in data.get("commands", []):
        base_cmd = cmd.strip().upper().replace('?', '')
        MASTER_CMD_MAP[base_cmd] = {"original": cmd.strip(), "group": group_name}
        if not cmd.strip().endswith('?'):
            MASTER_CMD_MAP[base_cmd + '?'] = {"original": cmd.strip() + '?', "group": group_name}

print(f"Loaded {len(MASTER_CMD_MAP)} command variants from mapping file.")

SECTION_MAP = {
    "GROUP": "group", "SYNTAX": "syntax", "ARGUMENTS": "arguments",
    "EXAMPLES": "examples", "RELATED COMMANDS": "related", "RELATED": "related",
    "RETURNS": "returns", "CONDITIONS": "conditions"
}


def expand_wildcards(option_list):
    expanded = []
    for opt in option_list:
        opt = opt.strip()
        if not opt:
            continue
        if re.search(r'CH<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"CH{i}" for i in range(1, 9)])
        elif re.search(r'MATH<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"MATH{i}" for i in range(1, 5)])
        elif re.search(r'REF<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"REF{i}" for i in range(1, 5)])
        elif re.search(r'BUS<[xn]>|B<[xn]>', opt, re.IGNORECASE):
            prefix = "B" if "B<" in opt else "BUS"
            expanded.extend([f"{prefix}{i}" for i in range(1, 9)])
        elif re.search(r'MEAS<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"MEAS{i}" for i in range(1, 9)])
        elif re.search(r'SEARCH<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"SEARCH{i}" for i in range(1, 9)])
        elif re.search(r'PLOT<[xn]>', opt, re.IGNORECASE):
            expanded.extend([f"PLOT{i}" for i in range(1, 5)])
        elif re.search(r'<NR[123]>|<QString>', opt, re.IGNORECASE):
            continue
        else:
            expanded.append(opt)
    seen = set()
    return [x for x in expanded if not (x in seen or seen.add(x))][:50]


def detect_params(command_header, syntax_list, arguments_text, examples=None):
    params = []
    
    # Extract default value from first SET example (not query)
    example_default = None
    if examples:
        for ex in examples:
            ex_text = ex.get('scpi', '') if isinstance(ex, dict) else str(ex)
            if '?' in ex_text:
                continue  # Skip queries
            
            # Example format: "COMMAND ARG description_starting_with_verb"
            # Find where description starts (verbs like sets, indicates, etc.)
            desc_match = re.search(r'\s+(sets|indicates|specifies|returns|turns|enables|queries|is|might|this)\s+', 
                                   ex_text, re.IGNORECASE)
            if desc_match:
                ex_text = ex_text[:desc_match.start()]  # Keep only SCPI part
            
            # Now extract argument (last token after space)
            if ' ' in ex_text:
                parts = ex_text.rsplit(' ', 1)
                if len(parts) == 2:
                    arg = parts[1].strip()
                    # Validate it looks like a value (number, ON/OFF, quoted string, CH1, etc.)
                    if re.match(r'^-?[\d.]+[eE]?[-+]?\d*$', arg) or \
                       re.match(r'^(ON|OFF|0|1)$', arg, re.IGNORECASE) or \
                       re.match(r'^".*"$', arg) or \
                       re.match(r'^[A-Z]+\d*$', arg, re.IGNORECASE):
                        example_default = arg
                        break
    
    # Pattern 0: Detect mnemonic placeholders in header (e.g., POWer<x>, CH<x>, MEAS<x>)
    # These are editable parts of the command path itself
    header_placeholders = re.findall(r'([A-Z]+)<([xn])>', command_header, re.IGNORECASE)
    for prefix, placeholder in header_placeholders:
        prefix_upper = prefix.upper()
        # Determine the param name and range based on prefix
        if prefix_upper in ['POWER', 'POWER']:
            params.append({"name": "power", "type": "integer", "required": True, 
                          "default": 1, "min": 1, "max": 8, "description": "Power measurement number (1-8)"})
        elif prefix_upper == 'CH':
            params.append({"name": "channel", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Channel number (1-8)"})
        elif prefix_upper == 'MATH':
            params.append({"name": "math", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Math waveform number (1-4)"})
        elif prefix_upper == 'REF':
            params.append({"name": "ref", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Reference waveform number (1-4)"})
        elif prefix_upper == 'MEAS':
            params.append({"name": "meas", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Measurement number (1-8)"})
        elif prefix_upper == 'SEARCH':
            params.append({"name": "search", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Search number (1-8)"})
        elif prefix_upper in ['B', 'BUS']:
            params.append({"name": "bus", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Bus number (1-8)"})
        elif prefix_upper == 'PLOT':
            params.append({"name": "plot", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Plot number (1-4)"})
        elif prefix_upper == 'HISTOGRAM':
            params.append({"name": "histogram", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Histogram number (1-4)"})
        elif prefix_upper == 'CURSOR':
            params.append({"name": "cursor", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 2, "description": "Cursor number (1-2)"})
        elif prefix_upper == 'CALLOUT':
            params.append({"name": "callout", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 8, "description": "Callout number (1-8)"})
        elif prefix_upper == 'MASK':
            params.append({"name": "mask", "type": "integer", "required": True,
                          "default": 1, "min": 1, "max": 4, "description": "Mask number (1-4)"})
    
    if syntax_list:
        full_syntax = " ".join(syntax_list)
        
        # Pattern 1: Look for trailing parameters NOT in braces (e.g., "COMMAND CH<x>")
        # Check each syntax entry individually for trailing source params
        for syn in syntax_list:
            if '?' in syn:
                continue  # Skip query syntax
            trailing_match = re.search(r'\s+(CH<[xn]>|MATH<[xn]>|REF<[xn]>)\s*$', syn, re.IGNORECASE)
            if trailing_match:
                trailing_param = trailing_match.group(1).upper()
                expanded = expand_wildcards([trailing_param])
                if expanded:
                    params.append({
                        "name": "source", "type": "enumeration", "required": True,
                        "options": expanded, "default": expanded[0],
                        "description": f"Options: {', '.join(expanded[:5])}..."
                    })
                break  # Only add once
        
        # Pattern 2: Enum in braces {ON|OFF|...}
        enum_matches = re.findall(r'\{([^}]+)\}', full_syntax)
        for match in enum_matches:
            raw_options = [o.strip() for o in match.split('|') if o.strip()]
            # Check if ONLY numeric placeholders (skip - handle later)
            if all(re.match(r'<NR[123]>|<QString>|<Block>', o, re.IGNORECASE) for o in raw_options):
                continue
            # Check if mixed: {ON|OFF|<NR1>} - means ONE param with multiple valid formats
            has_nr1 = any(re.match(r'<NR1>', o, re.IGNORECASE) for o in raw_options)
            # Filter out placeholders, keep only enum values
            enum_only = [o for o in raw_options if not re.match(r'<NR[123]>|<QString>|<Block>', o, re.IGNORECASE)]
            expanded_options = expand_wildcards(enum_only)
            if not expanded_options:
                continue
            param_name = "value"
            if "ON" in expanded_options and "OFF" in expanded_options:
                param_name = "state"
            elif any("CH" in o for o in expanded_options[:5]):
                param_name = "source"
            desc_preview = ", ".join(expanded_options[:5]) + ("..." if len(expanded_options) > 5 else "")
            if has_nr1:
                desc_preview += " (or 0/1)"
            params.append({
                "name": param_name, "type": "enumeration", "required": True,
                "options": expanded_options, "default": expanded_options[0] if expanded_options else None,
                "description": f"Options: {desc_preview}"
            })
            has_enum_arg = True
        
        # Add numeric/string params if no ENUM argument was found
        # (mnemonic params like bus/power don't count - they're for the path, not arguments)
        has_enum_arg = any(p.get('type') == 'enumeration' and p.get('name') not in ['bus', 'power', 'channel', 'math', 'ref', 'meas', 'search', 'plot', 'cursor', 'mask', 'histogram', 'callout', 'source'] for p in params)
        if not has_enum_arg:
            if re.search(r'<NR1>', full_syntax, re.IGNORECASE):
                # Try to get integer default from example
                int_default = None
                if example_default:
                    try:
                        int_default = int(float(example_default))
                    except:
                        pass
                params.append({"name": "value", "type": "integer", "required": True, 
                              "default": int_default, "description": "Integer value"})
            elif re.search(r'<NR[23]>', full_syntax, re.IGNORECASE):
                # Try to get float default from example
                float_default = None
                if example_default:
                    try:
                        float_default = float(example_default)
                    except:
                        pass
                params.append({"name": "value", "type": "float", "required": True,
                              "default": float_default, "description": "Floating point value"})
            elif re.search(r'<QString>', full_syntax, re.IGNORECASE):
                # Use example string as default (strip quotes)
                str_default = None
                if example_default:
                    str_default = example_default.strip('"\'')
                params.append({"name": "label", "type": "string", "required": True,
                              "default": str_default, "description": "Quoted string value"})
    if not params and arguments_text:
        lower_args = arguments_text.lower()
        if "integer" in lower_args:
            params.append({"name": "value", "type": "integer", "required": True})
        elif "float" in lower_args or "nr2" in lower_args or "nr3" in lower_args:
            params.append({"name": "value", "type": "float", "required": True})
        elif "string" in lower_args or "quoted" in lower_args:
            params.append({"name": "label", "type": "string", "required": True})
    return params


def detect_command_type(syntax_list, command_header):
    if not syntax_list:
        if command_header.endswith('?'):
            return {"commandType": "query", "hasQuery": True, "hasSet": False, "querySyntax": command_header, "setSyntax": None}
        else:
            return {"commandType": "both", "hasQuery": True, "hasSet": True, "querySyntax": command_header + "?", "setSyntax": command_header}
    full_syntax = " ".join(syntax_list)
    has_query = '?' in full_syntax
    # Check for SET: has value placeholder OR trailing source parameter (CH<x>, MATH<x>, etc.)
    has_set = bool(re.search(r'\{[^}]+\}|<NR[123]>|<QString>|<Block>|\s+CH<[xn]>|\s+MATH<[xn]>|\s+REF<[xn]>', full_syntax, re.IGNORECASE))
    if has_query and has_set:
        cmd_type = "both"
    elif has_query:
        cmd_type = "query"
    else:
        cmd_type = "set"
    query_syntax = None
    set_syntax = None
    for syn in syntax_list:
        syn = syn.strip()
        if '?' in syn and not query_syntax:
            query_syntax = syn
        elif '?' not in syn and not set_syntax:
            set_syntax = syn
    if not query_syntax and has_query:
        query_syntax = command_header + "?" if not command_header.endswith('?') else command_header
    if not set_syntax and has_set:
        set_syntax = command_header.replace('?', '')
    return {"commandType": cmd_type, "hasQuery": has_query, "hasSet": has_set, "querySyntax": query_syntax, "setSyntax": set_syntax}


def find_document():
    for full_path in POSSIBLE_FILES:
        if os.path.exists(full_path):
            return full_path
        return None
    

def clean_text(lines):
    if not lines:
        return None
    text = " ".join(lines).strip()
    return re.sub(r'\s+', ' ', text)


def is_bold(para):
    if para.style and para.style.font and para.style.font.bold:
        return True
    for run in para.runs:
        if run.bold:
            return True
    return False


def extract_commands(file_path):
    print(f"Loading document: {os.path.basename(file_path)}...")
    doc = docx.Document(file_path)
    commands = []
    current_cmd = None
    current_section = "description"
    buffer = []
    total_paras = len(doc.paragraphs)
    print(f"Scanning {total_paras} paragraphs...")

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        norm_text = text.upper().replace('?', '')
        is_header = norm_text in MASTER_CMD_MAP and is_bold(para)
        
        if is_header:
            if current_cmd:
                if buffer:
                    current_cmd[current_section].extend(buffer)
                commands.append(current_cmd)
            master_info = MASTER_CMD_MAP[norm_text]
            current_cmd = {
                "header": master_info["original"], "mapped_group": master_info["group"],
                "description": [], "group": [], "syntax": [], "arguments": [],
                "examples": [], "related": [], "returns": [], "conditions": [], "notes": []
            }
            current_section = "description"
            buffer = []
            continue
        
        if current_cmd:
            first_word_raw = text.split()[0] if text.split() else ""
            first_word = first_word_raw.upper().replace(':', '')
            if text.upper().startswith("RELATED COMMANDS"):
                first_word = "RELATED COMMANDS"
            if first_word in SECTION_MAP:
                if buffer:
                    current_cmd[current_section].extend(buffer)
                current_section = SECTION_MAP[first_word]
                buffer = []
                content_after = text[len(first_word_raw):].strip()
                if content_after.startswith(":"):
                    content_after = content_after[1:].strip()
                if content_after:
                    buffer.append(content_after)
                continue
            if text.upper().startswith("NOTE:") or text.upper().startswith("NOTE "):
                current_cmd["notes"].append(text)
                continue
            buffer.append(text)
        
        if i % 5000 == 0 and i > 0:
            print(f"  Progress: {i}/{total_paras} ({100*i/total_paras:.1f}%) - {len(commands)} commands", flush=True)

    if current_cmd:
        if buffer:
            current_cmd[current_section].extend(buffer)
        commands.append(current_cmd)
    return commands


def validate_syntax(syntax_list, command_header):
    base_cmd = command_header.replace('?', '').upper()
    prefix = base_cmd.split(':')[0] if ':' in base_cmd else base_cmd
    valid = []
    for syn in syntax_list:
        syn = syn.strip()
        if syn and syn.upper().startswith(prefix):
            valid.append(syn)
    return valid


def validate_examples(example_lines, command_header):
    prefix = command_header.replace('?', '').upper().split(':')[0]
    valid = []
    for line in example_lines:
        line = line.strip()
        if not line or not line[0].isupper():
            continue
        # Reject description text (common starting words)
        line_lower = line.lower()
        if line_lower.startswith(('this ', 'the ', 'a ', 'an ', 'note', 'see ', 'use ', 'for ', 
                                   'when ', 'if ', 'to ', 'in ', 'on ', 'it ', 'you ', 'requires')):
            continue
        # Must start with the command prefix and contain a colon (SCPI format)
        line_upper = line.upper()
        first_word = line.split()[0] if line.split() else ""
        if first_word.upper().startswith(prefix) and ':' in line:
            valid.append(line)
    return valid


def post_process(raw_commands):
    groups_dict = {}
    for cmd in raw_commands:
        header = cmd["header"]
        group_name = cmd["mapped_group"] or "Uncategorized"
        desc = clean_text(cmd["description"]) or ""
        args_text = clean_text(cmd["arguments"])
        validated_syntax = validate_syntax(cmd["syntax"], header)
        validated_examples = validate_examples(cmd["examples"], header)
        cmd_type_info = detect_command_type(validated_syntax, header)
        detected_params = detect_params(header, validated_syntax, args_text, validated_examples)
        
        processed_examples = []
        for line in validated_examples:
            # Pattern: SCPI_COMMAND [ARGUMENT] description_starting_with_verb
            # Description typically starts with: sets, queries, returns, indicates, specifies, turns, enables, might
            desc_start_pattern = r'\s+(sets|queries|returns|indicates|specifies|turns|enables|disables|might|is|will|this)\s+'
            match = re.search(desc_start_pattern, line, re.IGNORECASE)
            if match:
                code = line[:match.start()].strip()
                ex_desc = line[match.start():].strip()
            elif "might return" in line.lower():
                parts = line.split("might return", 1)
                code = parts[0].strip()
                ex_desc = "might return " + parts[1].strip() if len(parts) > 1 else ""
            else:
                # Fallback: split at first lowercase word (but keep uppercase args)
                parts = re.split(r'\s+(?=[a-z])', line, maxsplit=1)
                code = parts[0].strip() if parts else line
                ex_desc = parts[1].strip() if len(parts) > 1 else ""
            processed_examples.append({"scpi": code, "description": ex_desc, "codeExamples": {"scpi": {"code": code}}})
        
        name = header.split(':')[-1].replace('?', '').replace('<x>', '').replace('<n>', '')
        short_desc = desc[:100] + "..." if len(desc) > 100 else desc
        
        manual_entry = {
            "command": header, "header": header.split(':')[0],
            "mnemonics": header.replace('?', '').split(':'),
            "commandType": cmd_type_info["commandType"],
            "hasQuery": cmd_type_info["hasQuery"], "hasSet": cmd_type_info["hasSet"],
            "description": desc, "shortDescription": short_desc, "arguments": args_text,
            "examples": processed_examples,
            "relatedCommands": [c.strip() for c in clean_text(cmd["related"]).split()] if cmd["related"] else [],
            "commandGroup": group_name, "syntaxList": validated_syntax,
            "syntax": {"set": cmd_type_info["setSyntax"], "query": cmd_type_info["querySyntax"]},
            "manualReference": {"section": group_name}, "notes": cmd["notes"]
        }

        final_cmd_obj = {
            "scpi": header, "name": name, "description": desc, "shortDescription": short_desc,
            "group": group_name, "syntax": validated_syntax, "arguments": args_text,
            "params": detected_params, "examples": processed_examples,
            "relatedCommands": manual_entry["relatedCommands"],
            "conditions": clean_text(cmd["conditions"]), "returns": clean_text(cmd["returns"]),
            "notes": cmd["notes"], "example": processed_examples[0]["scpi"] if processed_examples else None,
            "commandType": cmd_type_info["commandType"],
            "hasQuery": cmd_type_info["hasQuery"], "hasSet": cmd_type_info["hasSet"],
            "_manualEntry": manual_entry
        }

        if group_name not in groups_dict:
            groups_dict[group_name] = {
                "name": group_name,
                "description": COMMAND_GROUPS.get(group_name, {}).get("description", ""),
                "commands": []
            }
        groups_dict[group_name]["commands"].append(final_cmd_obj)

    return {
        "version": "2.0", "manual": "4/5/6 Series MSO Programmer Manual",
        "groups": groups_dict,
        "metadata": {"total_commands": len(raw_commands), "total_groups": len(groups_dict)}
    }


if __name__ == "__main__":
    found_file = find_document()
    if not found_file:
        print("ERROR: Word document not found.")
        sys.exit(1)
    
    print("Extracting (Golden Key Verification Mode)...")
    raw_data = extract_commands(found_file)
    print(f"\nExtracted {len(raw_data)} commands matched against Mapping File.")
    
    final_json = post_process(raw_data)
    os.makedirs(os.path.dirname(OUTPUT_FILENAME), exist_ok=True)
    
    with open(OUTPUT_FILENAME, "w", encoding="utf-8") as f:
        json.dump(final_json, f, indent=2, ensure_ascii=False)
    
    print(f"\nSUCCESS! Output saved to: {OUTPUT_FILENAME}")
    print(f"  Commands: {final_json['metadata']['total_commands']}")
    print(f"  Groups: {final_json['metadata']['total_groups']}")

