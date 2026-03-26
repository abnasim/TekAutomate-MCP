"""
Patched extraction with debug logging for MEASUrement:CH<x>:REFLevels:ABSolute:TYPE
"""
import json
import re
import sys
import os

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from command_groups_mapping import COMMAND_GROUPS

# Build command lookup
ALL_COMMANDS = []
COMMAND_TO_GROUP = {}
for group_name, group_data in COMMAND_GROUPS.items():
    for cmd in group_data.get("commands", []):
        ALL_COMMANDS.append(cmd)
        COMMAND_TO_GROUP[cmd] = group_name

def get_font_name(run):
    if run.font and run.font.name:
        return run.font.name
    if run.style and hasattr(run.style, 'font') and run.style.font and run.style.font.name:
        return run.style.font.name
    return None

def is_tahoma(run):
    font_name = get_font_name(run)
    if font_name:
        return "tahoma" in font_name.lower()
    return False

def is_courier_new(run):
    font_name = get_font_name(run)
    if font_name:
        font_lower = font_name.lower()
        return "courier" in font_lower and "new" in font_lower
    return False

def extract_text_by_font(paragraph, font_check_func):
    parts = []
    for run in paragraph.runs:
        if font_check_func(run):
            parts.append(run.text.strip())
    return " ".join(parts) if parts else None

def extract_tahoma_text(paragraph):
    return extract_text_by_font(paragraph, is_tahoma)

def extract_courier_new_text(paragraph):
    return extract_text_by_font(paragraph, is_courier_new)

def is_command_in_master_list(text):
    if not text:
        return None
    text_upper = text.upper().replace('?', '')
    for cmd in ALL_COMMANDS:
        cmd_upper = cmd.upper().replace('?', '')
        text_norm = re.sub(r'<[xn]>', '<x>', text_upper)
        cmd_norm = re.sub(r'<[xn]>', '<x>', cmd_upper)
        if text_norm == cmd_norm:
            return cmd
    return None

def clean_text(text_or_list):
    if isinstance(text_or_list, list):
        return ' '.join(text_or_list).strip()
    return text_or_list.strip() if text_or_list else ""

# Load document
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
DOCX_PATH = os.path.join(PROJECT_ROOT, "4-5-6_MSO_Programmer_077189801_RevA (1).docx")

print(f"Loading document...")
doc = Document(DOCX_PATH)
print(f"Loaded. {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# Target command for debugging
DEBUG_CMD = "MEASUrement:CH<x>:REFLevels:ABSolute:TYPE"

commands = {}
current_cmd = None
state = "SEARCHING"
buffer = []

def log_debug(msg):
    """Log if we're processing our target command"""
    if current_cmd and DEBUG_CMD in current_cmd.get("scpi", ""):
        print(f"  [DEBUG {DEBUG_CMD}] {msg}")

def save_current_section():
    global buffer
    if not current_cmd:
        return
    
    text_content = clean_text(buffer)
    log_debug(f"save_current_section state={state}, buffer_len={len(buffer)}, text={text_content[:50] if text_content else 'empty'}...")
    
    if state == "DESCRIPTION":
        current_cmd["description"] = text_content
    elif state == "CONDITIONS":
        current_cmd["conditions"] = text_content
    elif state == "GROUP":
        current_cmd["group"] = text_content.strip()
    elif state == "SYNTAX":
        syntax_lines = [line.strip() for line in buffer if line.strip()]
        if syntax_lines:
            if not current_cmd.get("syntax"):
                current_cmd["syntax"] = []
            current_cmd["syntax"].extend(syntax_lines)
            log_debug(f"Extended syntax from buffer: {syntax_lines}")
    elif state == "ARGUMENTS":
        current_cmd["arguments"] = text_content
    elif state == "EXAMPLES":
        # Parse examples from buffer
        if buffer:
            if not current_cmd.get("examples"):
                current_cmd["examples"] = []
            
            for line in buffer:
                line = line.strip()
                if not line:
                    continue
                
                # Pattern: "CMD? might return VALUE, description"
                if "might return" in line.lower():
                    parts = re.split(r'\s+might\s+return\s+', line, maxsplit=1, flags=re.IGNORECASE)
                    if len(parts) == 2:
                        scpi_part = parts[0].strip() + " might return"
                        after_return = parts[1].strip()
                        match = re.match(r'^([A-Z0-9:,\-_\s]+?)(?:\s*,\s*|\s+)(indicating|that|which|the|a)\s', after_return, re.IGNORECASE)
                        if match:
                            returned_val = match.group(1).strip()
                            desc_start = match.start(2)
                            description = after_return[desc_start:].strip()
                            scpi_part = f"{parts[0].strip()} might return {returned_val}"
                        else:
                            description = after_return
                        current_cmd["examples"].append({"scpi": scpi_part, "description": description})
                        continue
                
                # Standard: "CMD VALUE sets/queries description"
                match = re.match(r'^(.+?)\s+(sets|queries|returns|indicates|turns|specifies|enables|disables)\s+(.+)$', line, re.IGNORECASE)
                if match:
                    scpi_part = match.group(1).strip()
                    verb = match.group(2).lower()
                    desc_rest = match.group(3).strip()
                    description = f"{verb} {desc_rest}"
                    current_cmd["examples"].append({"scpi": scpi_part, "description": description})
                else:
                    # Fallback
                    match = re.match(r'^([A-Z0-9:<>?\[\]{}\s\-_,]+?)\s+([a-z].*)$', line)
                    if match:
                        current_cmd["examples"].append({"scpi": match.group(1).strip(), "description": match.group(2).strip()})
                    else:
                        current_cmd["examples"].append({"scpi": line, "description": ""})
    
    buffer = []

def finalize_command():
    global current_cmd, state, buffer
    
    if not current_cmd:
        return
    
    log_debug(f"finalize_command() called")
    log_debug(f"  syntax before save: {current_cmd.get('syntax', [])}")
    log_debug(f"  conditions before save: {current_cmd.get('conditions')}")
    log_debug(f"  arguments before save: {current_cmd.get('arguments')}")
    
    save_current_section()
    
    log_debug(f"  syntax after save: {current_cmd.get('syntax', [])}")
    
    # Deduplication
    cmd_key = current_cmd["scpi"]
    if cmd_key in commands:
        existing = commands[cmd_key]
        log_debug(f"  DEDUP: existing entry found!")
        log_debug(f"  existing syntax: {existing.get('syntax', [])}")
        log_debug(f"  current syntax: {current_cmd.get('syntax', [])}")
        
        for key in ["description", "conditions", "arguments", "examples", "group"]:
            if not existing.get(key) and current_cmd.get(key):
                existing[key] = current_cmd[key]
                log_debug(f"  Copied {key} to existing")
        
        if current_cmd.get("syntax"):
            existing["syntax"].extend([s for s in current_cmd["syntax"] if s not in existing.get("syntax", [])])
            log_debug(f"  Extended existing syntax: {existing['syntax']}")
    else:
        commands[cmd_key] = current_cmd
        log_debug(f"  Stored as new entry")
    
    current_cmd = None
    state = "SEARCHING"
    buffer = []

# Process paragraphs
print("Processing paragraphs...")
for para_num, paragraph in enumerate(doc.paragraphs):
    line = paragraph.text.strip()
    if not line:
        continue
    
    line_lower = line.lower()
    
    # Keyword checks
    if line_lower.startswith("conditions"):
        save_current_section()
        state = "CONDITIONS"
        buffer = []
        continue
    elif line_lower.startswith("group"):
        save_current_section()
        state = "GROUP"
        buffer = []
        continue
    elif line_lower.startswith("syntax"):
        log_debug(f"[Para {para_num}] 'Syntax' keyword detected")
        save_current_section()
        state = "SYNTAX"
        buffer = []
        continue
    elif line_lower.startswith("arguments"):
        log_debug(f"[Para {para_num}] 'Arguments' keyword detected")
        save_current_section()
        state = "ARGUMENTS"
        buffer = []
        continue
    elif line_lower.startswith("examples"):
        save_current_section()
        state = "EXAMPLES"
        buffer = []
        continue
    
    # Check for Tahoma command header
    tahoma_text = extract_tahoma_text(paragraph)
    if tahoma_text:
        first_word = tahoma_text.split()[0] if tahoma_text.split() else ""
        matched_cmd = is_command_in_master_list(first_word)
        if matched_cmd:
            if current_cmd and DEBUG_CMD in current_cmd.get("scpi", ""):
                print(f"  [DEBUG] New command header found at para {para_num}: {matched_cmd}")
                print(f"  [DEBUG] Finalizing {DEBUG_CMD}...")
            
            finalize_command()
            current_cmd = {
                "scpi": matched_cmd,
                "description": None,
                "conditions": None,
                "group": COMMAND_TO_GROUP.get(matched_cmd),
                "syntax": [],
                "arguments": None,
                "examples": None,
            }
            state = "DESCRIPTION"
            buffer = []
            
            if DEBUG_CMD in matched_cmd:
                print(f"\n[DEBUG] Started processing {DEBUG_CMD} at paragraph {para_num}")
            continue
    
    # Check for Courier New syntax
    if state == "SYNTAX" and current_cmd:
        courier_text = extract_courier_new_text(paragraph)
        if courier_text:
            if not current_cmd.get("syntax"):
                current_cmd["syntax"] = []
            current_cmd["syntax"].append(courier_text.strip())
            log_debug(f"[Para {para_num}] Captured syntax via Courier: {courier_text[:50]}...")
            continue
    
    # Buffer accumulation
    if state != "SEARCHING" and current_cmd:
        buffer.append(line)
        if DEBUG_CMD in current_cmd.get("scpi", ""):
            log_debug(f"[Para {para_num}] Buffered: {line[:50]}...")

# Finalize last command
finalize_command()

print(f"\nProcessed {len(commands)} commands")

# Check our target command
if DEBUG_CMD in commands:
    print(f"\n{'='*60}")
    print(f"RESULT FOR {DEBUG_CMD}:")
    print(f"{'='*60}")
    cmd = commands[DEBUG_CMD]
    print(json.dumps(cmd, indent=2, default=str))
else:
    print(f"\n{DEBUG_CMD} NOT FOUND in extracted commands!")

