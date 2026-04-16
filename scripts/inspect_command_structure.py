"""Inspect actual command structure in Word document"""
from docx import Document
import re
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
DOCX_PATH = os.path.join(PROJECT_ROOT, "4-5-6_MSO_Programmer_077189801_RevA.docx")

doc = Document(DOCX_PATH)
CMD_PATTERN = re.compile(r'^([*][A-Za-z]{2,}\??)$|^([A-Za-z][A-Za-z0-9<>]*:[A-Za-z0-9<>:]+(?:\?)?)$')

print("=== Finding ACQuire:STATE command (known to exist) ===")
found = False

for para_num, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    words = text.split()
    if words and words[0] == "ACQuire:STATE":
        found = True
        print(f"\nFound at paragraph {para_num}")
        print(f"Full text: {text[:300]}")
        
        # Get surrounding paragraphs for context
        start = max(0, para_num - 3)
        end = min(len(doc.paragraphs), para_num + 25)
        print(f"\n=== Context (paragraphs {start} to {end}) ===")
        for i in range(start, end):
            ptext = doc.paragraphs[i].text.strip()
            if ptext:
                # Mark the command line
                marker = ">>> " if i == para_num else "    "
                print(f"{marker}[{i}] {ptext[:120]}")
        break

# Search for command definitions (commands followed by descriptions)
print("\n=== Searching for command definitions with descriptions ===")
count = 0
for para_num, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    words = text.split()
    if words:
        first_word = words[0]
        # Look for commands that have description on same line or next line
        if CMD_PATTERN.match(first_word) and len(words) > 1:
            # Check if second word is not another command
            second_part = " ".join(words[1:])
            if not CMD_PATTERN.match(second_part.split()[0] if second_part.split() else ""):
                count += 1
                if count <= 5:
                    print(f"\n[{para_num}] Command: {first_word}")
                    print(f"    Description: {second_part[:150]}")
                    # Show next few paragraphs
                    for i in range(para_num + 1, min(len(doc.paragraphs), para_num + 8)):
                        ptext = doc.paragraphs[i].text.strip()
                        if ptext:
                            # Check for section headers
                            if any(ptext.lower().startswith(h) for h in ["syntax", "arguments", "examples", "group", "related"]):
                                print(f"    [{i}] {ptext[:100]}")
                            elif len(ptext) < 150:  # Short lines might be section content
                                print(f"    [{i}] {ptext[:100]}")
                    print("---")
                if count >= 5:
                    break

# Search for actual command definitions with section structure
print("\n=== Searching for command definitions with section headers ===")
section_keywords = ["Syntax", "Arguments", "Examples", "Group", "Related Commands", "Description", "Returns", "Conditions"]
found_definitions = []

for para_num, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    words = text.split()
    
    # Look for a command header
    if words and CMD_PATTERN.match(words[0]):
        # Check if following paragraphs have section headers
        has_sections = False
        section_found = []
        for i in range(para_num + 1, min(len(doc.paragraphs), para_num + 15)):
            next_text = doc.paragraphs[i].text.strip()
            if any(next_text.lower().startswith(kw.lower()) for kw in section_keywords):
                has_sections = True
                section_found.append(next_text.split()[0])
        
        if has_sections:
            found_definitions.append((para_num, words[0], section_found))
            if len(found_definitions) <= 3:
                print(f"\n[{para_num}] Command: {words[0]}")
                print(f"    Full line: {text[:150]}")
                print(f"    Sections found: {section_found}")
                # Show context
                for i in range(para_num, min(len(doc.paragraphs), para_num + 12)):
                    ptext = doc.paragraphs[i].text.strip()
                    if ptext:
                        marker = ">>> " if i == para_num else "    "
                        print(f"{marker}[{i}] {ptext[:120]}")
                print("---")
            
            if len(found_definitions) >= 3:
                break

print(f"\nFound {len(found_definitions)} command definitions with sections")

