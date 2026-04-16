"""
Fast extraction script - based on working debug version
Captures: syntax with arguments, examples, conditions, arguments text
"""
import json
import re
import sys
import os

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from command_groups_mapping import COMMAND_GROUPS

# Build command lookup
ALL_COMMANDS = []
COMMAND_TO_GROUP = {}
for group_name, group_data in COMMAND_GROUPS.items():
    for cmd in group_data.get("commands", []):
        ALL_COMMANDS.append(cmd)
        COMMAND_TO_GROUP[cmd] = group_name

print(f"Loaded {len(ALL_COMMANDS)} commands from {len(COMMAND_GROUPS)} groups")

def get_font_name(run):
    if run.font and run.font.name:
        return run.font.name
    if run.style and hasattr(run.style, 'font') and run.style.font and run.style.font.name:
        return run.style.font.name
    return None

def is_tahoma(run):
    font_name = get_font_name(run)
    return "tahoma" in font_name.lower() if font_name else False

def is_courier_new(run):
    font_name = get_font_name(run)
    if font_name:
        return "courier" in font_name.lower() and "new" in font_name.lower()
    return False

def extract_courier_new_text(paragraph):
    parts = []
    for run in paragraph.runs:
        if is_courier_new(run):
            parts.append(run.text.strip())
    return " ".join(parts) if parts else None

def extract_tahoma_text(paragraph):
    parts = []
    for run in paragraph.runs:
        if is_tahoma(run):
            parts.append(run.text.strip())
    return " ".join(parts) if parts else None

def is_command_in_master_list(text):
    if not text:
        return None
    text_upper = text.upper().strip()
    # Direct match
    if text_upper in [c.upper() for c in ALL_COMMANDS]:
        for c in ALL_COMMANDS:
            if c.upper() == text_upper:
                return c
    # Placeholder pattern match
    pattern_text = re.sub(r'\d+', '<x>', text_upper)
    for cmd in ALL_COMMANDS:
        cmd_pattern = cmd.upper()
        if pattern_text == cmd_pattern:
            return cmd
    return None

def clean_text(buffer):
    return "\n".join(line.strip() for line in buffer if line.strip())

# Main extraction
script_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(script_dir)
docx_files = [f for f in os.listdir(parent_dir) if f.endswith('.docx') and not f.startswith('~')]
docx_path = os.path.join(parent_dir, docx_files[0]) if docx_files else None

if not docx_path:
    print("ERROR: No Word document found")
    sys.exit(1)

print(f"Loading: {docx_path}")
doc = Document(docx_path)
print(f"Loaded. {len(doc.paragraphs)} paragraphs")

commands = {}
current_cmd = None
state = "SEARCHING"
buffer = []

def save_current_section():
    global buffer
    if not current_cmd:
        return
    
    text_content = clean_text(buffer)
    
    if state == "DESCRIPTION":
        current_cmd["description"] = text_content
    elif state == "CONDITIONS":
        current_cmd["conditions"] = text_content
    elif state == "GROUP":
        # Only take first line, ignore any syntax that follows
        first_line = buffer[0].strip() if buffer else ""
        # Map to known group from master list, or use extracted
        mapped = COMMAND_TO_GROUP.get(current_cmd["scpi"])
        current_cmd["group"] = mapped or first_line
    elif state == "SYNTAX":
        syntax_lines = [line.strip() for line in buffer if line.strip()]
        if syntax_lines:
            if not current_cmd.get("syntax"):
                current_cmd["syntax"] = []
            current_cmd["syntax"].extend(syntax_lines)
    elif state == "ARGUMENTS":
        current_cmd["arguments"] = text_content
    elif state == "EXAMPLES":
        if buffer:
            if not current_cmd.get("examples"):
                current_cmd["examples"] = []
            for line in buffer:
                line = line.strip()
                if not line:
                    continue
                # "might return" pattern
                if "might return" in line.lower():
                    parts = re.split(r'\s+might\s+return\s+', line, maxsplit=1, flags=re.IGNORECASE)
                    if len(parts) == 2:
                        scpi_part = parts[0].strip()
                        desc = parts[1].strip()
                        current_cmd["examples"].append({"scpi": scpi_part, "description": f"might return {desc}"})
                        continue
                # "sets/queries" pattern
                match = re.match(r'^(.+?)\s+(sets|queries|returns|indicates|specifies)\s+(.+)$', line, re.IGNORECASE)
                if match:
                    current_cmd["examples"].append({"scpi": match.group(1).strip(), "description": f"{match.group(2)} {match.group(3)}"})
                else:
                    match = re.match(r'^([A-Z0-9:<>?\s\-_]+?)\s+([a-z].*)$', line)
                    if match:
                        current_cmd["examples"].append({"scpi": match.group(1).strip(), "description": match.group(2).strip()})
                    else:
                        current_cmd["examples"].append({"scpi": line, "description": ""})
    elif state == "RETURNS":
        current_cmd["returns"] = text_content
    elif state == "RELATED":
        related = [cmd.strip() for cmd in text_content.split() if is_command_in_master_list(cmd.strip())]
        if related:
            current_cmd["relatedCommands"] = related
    
    buffer = []

def finalize_command():
    global current_cmd, state, buffer
    if not current_cmd:
        return
    
    save_current_section()
    
    # ALWAYS use master list group mapping (ignore extracted group - it's often malformed)
    current_cmd["group"] = COMMAND_TO_GROUP.get(current_cmd["scpi"], "Miscellaneous")
    
    # Normalize
    for key in ["relatedCommands", "conditions", "returns", "examples", "arguments"]:
        if not current_cmd.get(key):
            current_cmd[key] = None
    if not current_cmd.get("syntax"):
        current_cmd["syntax"] = []
    
    # Store
    cmd_key = current_cmd["scpi"]
    if cmd_key in commands:
        existing = commands[cmd_key]
        for key in ["description", "group", "arguments", "examples", "conditions", "returns"]:
            if not existing.get(key) and current_cmd.get(key):
                existing[key] = current_cmd[key]
        if current_cmd.get("syntax"):
            if not existing.get("syntax"):
                existing["syntax"] = []
            existing["syntax"].extend([s for s in current_cmd["syntax"] if s not in existing["syntax"]])
    else:
        commands[cmd_key] = current_cmd
    
    current_cmd = None
    state = "SEARCHING"
    buffer = []

print("Processing paragraphs...")
progress_interval = len(doc.paragraphs) // 20

for i, paragraph in enumerate(doc.paragraphs):
    if progress_interval > 0 and i % progress_interval == 0:
        print(f"  {i}/{len(doc.paragraphs)} ({100*i//len(doc.paragraphs)}%) - {len(commands)} commands", flush=True)
    
    line = paragraph.text.strip()
    if not line:
        continue
    
    line_lower = line.lower()
    
    # Notes
    if line_lower.startswith("note"):
        if current_cmd:
            if "notes" not in current_cmd:
                current_cmd["notes"] = []
            current_cmd["notes"].append(line)
        continue
    
    # Keywords
    if line_lower in ["conditions", "condition"]:
        save_current_section()
        state = "CONDITIONS"
        continue
    if line_lower == "group":
        save_current_section()
        state = "GROUP"
        continue
    if line_lower == "syntax":
        save_current_section()
        state = "SYNTAX"
        continue
    if line_lower in ["arguments", "argument"]:
        save_current_section()
        state = "ARGUMENTS"
        continue
    if line_lower in ["examples", "example"]:
        save_current_section()
        state = "EXAMPLES"
        continue
    if line_lower in ["returns", "return"]:
        save_current_section()
        state = "RETURNS"
        continue
    if "related commands" in line_lower or line_lower == "related":
        save_current_section()
        state = "RELATED"
        continue
    
    # Command header detection (Tahoma font)
    tahoma_text = extract_tahoma_text(paragraph)
    if tahoma_text:
        words = tahoma_text.split()
        if words:
            matched = is_command_in_master_list(words[0])
            if matched:
                finalize_command()
                current_cmd = {
                    "scpi": matched,
                    "description": None,
                    "conditions": None,
                    "group": COMMAND_TO_GROUP.get(matched),
                    "syntax": [],
                    "relatedCommands": None,
                    "arguments": None,
                    "examples": None,
                    "returns": None,
                    "notes": []
                }
                state = "DESCRIPTION"
                buffer = []
                continue
    
    # Syntax capture (Courier New font)
    if state == "SYNTAX" and current_cmd:
        courier_text = extract_courier_new_text(paragraph)
        if courier_text:
            if not current_cmd.get("syntax"):
                current_cmd["syntax"] = []
            current_cmd["syntax"].append(courier_text.strip())
            continue
    
    # Buffer content
    if current_cmd and state != "SEARCHING":
        buffer.append(line)

# Finalize last command
finalize_command()

print(f"\nExtracted {len(commands)} commands")

# Build output structure
groups_dict = {}
for cmd_key, cmd in commands.items():
    group_name = cmd.get("group", "Miscellaneous")
    if group_name not in groups_dict:
        groups_dict[group_name] = {"name": group_name, "description": "", "commands": []}
    
    # Build params from syntax
    params = []
    param_names = set()
    syntax_lines = cmd.get("syntax", [])
    
    # Extract <x> placeholders from command with meaningful names
    # Look for patterns like CH<x>, BUS<x>, MATH<x>, etc.
    # Official SCPI Mnemonic Patterns from Tektronix Programmer Manual
    # Each entry: (regex_pattern, param_name, description, min_value, max_value_or_None)
    # Mnemonic definitions with clearer descriptions showing allowed values
    MNEMONIC_DEFINITIONS = [
        # Bus mnemonics - Table 6
        (r'B<x>', 'bus', 'B<x> where x is the bus number (1-16)', 1, 16),
        
        # Channel mnemonics - Table 7
        (r'CH<x>_D<x>', 'channel', 'CH<x> where x is channel number (1-8)', 1, 8),
        (r'CH<x>', 'channel', 'CH<x> where x is channel number (1-8)', 1, 8),
        
        # Cursor mnemonics - Table 8
        (r'CURSOR<x>', 'cursor', 'CURSOR<x> where x is cursor number (1 or 2)', 1, 2),
        
        # Math specifier - Table 9
        (r'MATH<x>', 'math', 'MATH<x> where x is math waveform number (1-4)', 1, 4),
        
        # Measurement specifier - Table 10
        (r'MEAS<x>', 'measurement', 'MEAS<x> where x is measurement number (1-8)', 1, 8),
        
        # Reference waveform - Table 11
        (r'REF<x>_D<x>', 'reference', 'REF<x> where x is reference number (1-4)', 1, 4),
        (r'REF<x>', 'reference', 'REF<x> where x is reference waveform (1-4)', 1, 4),
        
        # View mnemonics - Table 12
        (r'WAVEView<x>', 'waveview', 'WAVEView<x> where x must be 1', 1, 1),
        (r'PLOTView<x>', 'plotview', 'PLOTView<x> where x must be 1', 1, 1),
        (r'MATHFFTView<x>', 'mathfftview', 'MATHFFTView<x> where x must be 1', 1, 1),
        
        # Search mnemonics - Table 13
        (r'SEARCH<x>', 'search', 'SEARCH<x> where x is search number (1-8)', 1, 8),
        
        # Zoom mnemonics - Table 14
        (r'ZOOM<x>', 'zoom', 'ZOOM<x> where x must be 1', 1, 1),
        
        # Power measurement
        (r'POWer<x>', 'power', 'POWer<x> where x is power measurement badge number (1-8)', 1, 8),
        
        # Histogram
        (r'HISTogram<x>', 'histogram', 'HISTogram<x> where x is histogram number (1-4)', 1, 4),
        
        # Plot
        (r'PLOT<x>', 'plot', 'PLOT<x> where x is plot number (1-4)', 1, 4),
        
        # Digital bit (D0-D7)
        (r'D<x>', 'digital_bit', 'D<x> where x is digital bit (0-7)', 0, 7),
        
        # Mask
        (r'MASK<x>', 'mask', 'MASK<x> where x is mask number (1-8)', 1, 8),
        
        # Other common patterns
        (r'CALLOUTS<x>', 'callout', 'CALLOUTS<x> where x is callout number', 1, 8),
        (r'ACTONEVent<x>', 'actonevent', 'ACTONEVent<x> where x is event number', 1, 8),
        (r'LICENSE<x>', 'license', 'LICENSE<x> where x is license number', 1, 8),
        (r'RAIL<x>', 'rail', 'RAIL<x> where x is rail number (1-7)', 1, 7),
        (r'SOURCE<x>', 'source_num', 'SOURCE<x> where x is source number', 1, 4),
    ]
    
    mnemonic_patterns = [(p[0], p[1], p[2]) for p in MNEMONIC_DEFINITIONS]
    
    for pattern, pname, desc in mnemonic_patterns:
        if re.search(pattern, cmd_key, re.IGNORECASE):
            if pname not in param_names:
                params.append({"name": pname, "type": "integer", "required": True, "default": 1, "description": desc})
                param_names.add(pname)
    
    # Parse {A|B} pattern from command header itself (e.g., TRIGger:{A|B}:BUS)
    header_enum_match = re.search(r'\{([A-Z|]+)\}', cmd_key, re.IGNORECASE)
    if header_enum_match:
        header_options = [o.strip() for o in header_enum_match.group(1).split('|') if o.strip()]
        if header_options and "trigger_type" not in param_names:
            params.append({
                "name": "trigger_type",
                "type": "enumeration",
                "required": True,
                "default": header_options[0],
                "options": header_options,
                "description": f"Trigger type: {' or '.join(header_options)}"
            })
            param_names.add("trigger_type")
    
    # Separate SET and QUERY syntax lines
    # First, extract set parts from combined syntax lines (e.g., "CMD {args} CMD?")
    set_syntax_lines = []
    query_syntax_lines = []
    
    for s in syntax_lines:
        s = s.strip()
        if not s:
            continue
        
        if "?" in s:
            # Check if this is a combined set+query syntax line
            # Look for pattern: COMMAND [args] COMMAND?
            header = cmd_key.split(" ")[0].split("?")[0]
            if header and s.count(header) >= 2:
                # Find the second occurrence which is the query
                first_pos = s.find(header)
                second_pos = s.find(header, first_pos + len(header))
                if second_pos > 0:
                    # Split into set and query parts
                    set_part = s[:second_pos].strip()
                    query_part = s[second_pos:].strip()
                    if set_part:
                        set_syntax_lines.append(set_part)
                    if query_part:
                        query_syntax_lines.append(query_part)
                    continue
            
            # Try regex pattern
            query_match = re.search(r'\s+([A-Za-z:]+(?:<x>)?[A-Za-z:]*\?)\s*$', s)
            if query_match:
                potential_set = s[:query_match.start()].strip()
                potential_query = query_match.group(1).strip()
                if potential_set and ('{' in potential_set or '<NR' in potential_set or '<QString' in potential_set):
                    set_syntax_lines.append(potential_set)
                    query_syntax_lines.append(potential_query)
                    continue
            
            # Pure query line
            query_syntax_lines.append(s)
        else:
            # Pure set line
            set_syntax_lines.append(s)
    
    # Extract standalone <NR1>, <NR2>, <NR3>, <QString> from SET syntax AND arguments only
    # Query commands don't have value parameters - they only return values
    set_syntax_text = " ".join(set_syntax_lines)
    args_text = cmd.get("arguments", "") or ""
    combined_text = set_syntax_text + " " + args_text
    
    # <NR1> = integer, <NR2> = float without exponent, <NR3> = float with exponent
    if re.search(r'<NR[123]>', combined_text):
        # Check if it's already part of a {} group
        if not re.search(r'\{[^}]*<NR[123]>[^}]*\}', combined_text):
            if "value" not in param_names:
                nr_type = "integer" if "<NR1>" in combined_text else "number"
                params.append({"name": "value", "type": nr_type, "required": True, "default": 1, "description": "Numeric value"})
                param_names.add("value")
    
    # <QString> = quoted string (check both syntax and arguments)
    if re.search(r'<QString>', combined_text, re.IGNORECASE):
        if not re.search(r'\{[^}]*<QString>[^}]*\}', combined_text, re.IGNORECASE):
            if "value" not in param_names:
                params.append({"name": "value", "type": "string", "required": True, "description": "Quoted string value"})
                param_names.add("value")
    
    # Extract {OPT1|OPT2} from SET syntax only (query commands don't have value parameters)
    for s in set_syntax_lines:
        for m in re.finditer(r'\{([^}]+)\}', s):
            raw_options = [o.strip() for o in m.group(1).split('|') if o.strip()]
            if not raw_options:
                continue
            
            # Skip if this is a simple {A|B} pattern already captured as trigger_type
            if all(len(opt) == 1 and opt.isupper() for opt in raw_options):
                continue  # Already handled by header_enum_match above
            
            # Expand mnemonic options like CH<x>, MATH<x>, REF<x>
            expanded_options = []
            has_mnemonic = False
            for opt in raw_options:
                if re.match(r'^CH<x>$', opt, re.IGNORECASE):
                    expanded_options.extend([f'CH{i}' for i in range(1, 9)])
                    has_mnemonic = True
                elif re.match(r'^MATH<x>$', opt, re.IGNORECASE):
                    expanded_options.extend([f'MATH{i}' for i in range(1, 5)])
                    has_mnemonic = True
                elif re.match(r'^REF<x>$', opt, re.IGNORECASE):
                    expanded_options.extend([f'REF{i}' for i in range(1, 5)])
                    has_mnemonic = True
                elif re.match(r'^B<x>$', opt, re.IGNORECASE):
                    expanded_options.extend([f'B{i}' for i in range(1, 9)])
                    has_mnemonic = True
                elif re.match(r'^MEAS<x>$', opt, re.IGNORECASE):
                    expanded_options.extend([f'MEAS{i}' for i in range(1, 9)])
                    has_mnemonic = True
                elif opt.startswith('<') and opt.endswith('>'):
                    # Placeholders like <file_path>, <NR1>, <QString> - add as special input option
                    placeholder_name = opt[1:-1].lower()
                    if 'path' in placeholder_name or 'file' in placeholder_name or 'string' in placeholder_name:
                        expanded_options.append('<custom>')  # Signal that custom input is allowed
                    elif 'nr' in placeholder_name.lower():
                        expanded_options.append('<number>')  # Signal numeric input allowed
                    continue
                else:
                    expanded_options.append(opt)
            
            if not expanded_options:
                # Only placeholders found - this is a text/number input, not enum
                placeholder = raw_options[0] if raw_options else "value"
                if "NR" in placeholder.upper():
                    if "value" not in param_names:
                        params.append({"name": "value", "type": "number", "required": True, "default": 1})
                        param_names.add("value")
                elif "string" in placeholder.lower() or "path" in placeholder.lower():
                    if "value" not in param_names:
                        params.append({"name": "value", "type": "string", "required": True})
                        param_names.add("value")
                continue
            
            if "value" not in param_names:
                param_name = "source" if has_mnemonic else "value"
                params.append({
                    "name": param_name,
                    "type": "enumeration",
                    "required": True,
                    "default": expanded_options[0],
                    "options": expanded_options,
                    "description": f"One of: {', '.join(expanded_options[:5])}{'...' if len(expanded_options) > 5 else ''}"
                })
                param_names.add("value")
    
    # If no value param yet, try to infer from SET examples only (skip query examples)
    if "value" not in param_names and cmd.get("examples"):
        for ex in cmd.get("examples", []):
            ex_scpi = ex.get("scpi", "")
            # Skip query examples (they don't have value arguments)
            if "?" in ex_scpi:
                continue
            # Check if example has a value argument (after the command path)
            # e.g., "POWer:POWer3:CLRESPONSE:CONSTAMPlitude 120"
            parts = ex_scpi.split()
            if len(parts) >= 2:
                value_part = parts[-1]
                # Check if it's a number
                if re.match(r'^-?\d+\.?\d*$', value_part):
                    params.append({"name": "value", "type": "number", "required": True, "default": 1, "description": "Numeric value"})
                    param_names.add("value")
                    break
                # Check if it's an enumeration option (uppercase word)
                elif re.match(r'^[A-Z]+\d*$', value_part):
                    # It's likely an enum option, but we don't know all options
                    # Skip for now as we can't infer the full option list
                    pass
    
    # If still no value param and description mentions set, add a generic value param
    # Only for SET commands, not query commands
    desc = cmd.get("description", "") or ""
    if "value" not in param_names and "sets" in desc.lower() and not cmd_key.endswith("?") and set_syntax_lines:
        # Check for numeric range in arguments
        args_text = cmd.get("arguments", "") or ""
        if re.search(r'<NR[123]>', args_text, re.IGNORECASE) or "number" in args_text.lower():
            params.append({"name": "value", "type": "number", "required": True, "default": 1, "description": "Numeric value"})
            param_names.add("value")
    
    # Short description
    short_desc = desc[:80] if len(desc) > 80 else desc
    
    # Name from command
    name_parts = cmd_key.split(":")
    name = name_parts[-1].replace("?", "").replace("<x>", "").replace("<", "").replace(">", "")
    name = name.capitalize() if name else "Command"
    
    # Build manualEntry
    # Check if this is a query-only command (description contains "query-only" or command ends with ? and has no set syntax)
    desc_lower = desc.lower() if desc else ""
    is_query_only = (
        "query-only" in desc_lower or 
        "(query only)" in desc_lower or 
        "(Query Only)" in desc or
        (cmd_key.endswith("?") and not any("?" not in s for s in syntax_lines if s.strip()))
    )
    
    set_syntax = ""
    query_syntax = ""
    
    for s in syntax_lines:
        s = s.strip()
        if not s:
            continue
            
        # Check if line contains BOTH set and query syntax (e.g., "CMD {args} CMD?")
        # Pattern: Contains "?" and also contains arguments like {}, <NR>, <QString>, etc.
        if "?" in s:
            # First, try the alternative pattern: Split on the command header that appears twice
            # This is more reliable than regex for complex command patterns
            # e.g., "BUS:B<x>:ARINC429A:SOUrce {CH<x>|MATH<x>|REF<x>} BUS:B<x>:ARINC429A:SOUrce?"
            header = cmd_key.split(" ")[0].split("?")[0]
            if header and s.count(header) >= 2:
                # Find the second occurrence which is the query
                first_pos = s.find(header)
                second_pos = s.find(header, first_pos + len(header))
                if second_pos > 0:
                    potential_set = s[:second_pos].strip()
                    potential_query = s[second_pos:].strip()
                    if potential_set and not set_syntax:
                        set_syntax = potential_set
                    if potential_query and not query_syntax:
                        query_syntax = potential_query
                    continue
            
            # Try to find if there's a SET command before the QUERY command using regex
            # Look for pattern: COMMAND [args] COMMAND?
            # The query command typically starts with the same mnemonic path and ends with ?
            # Improved regex to handle complex command patterns with <x> placeholders
            query_match = re.search(r'\s+([A-Za-z:]+(?:<x>)?[A-Za-z:]*\?)\s*$', s)
            if query_match:
                # Everything before the query command is the SET syntax
                potential_set = s[:query_match.start()].strip()
                potential_query = query_match.group(1).strip()
                
                # Validate: SET should have arguments, QUERY should just end with ?
                if potential_set and ('{' in potential_set or '<NR' in potential_set or '<QString' in potential_set):
                    if not set_syntax:
                        set_syntax = potential_set
                    if not query_syntax:
                        query_syntax = potential_query
                    continue
            
            # Simple case: just a query line
            if not query_syntax:
                query_syntax = s
        else:
            # No "?" means this is a SET syntax line
            if not set_syntax:
                set_syntax = s
    
    # For query-only commands, don't create a set syntax
    if is_query_only:
        set_syntax = None
        if not query_syntax:
            query_syntax = cmd_key if cmd_key.endswith("?") else cmd_key + "?"
    else:
        # For non-query-only commands, ensure both syntaxes exist if needed
        # Only set default set_syntax if we don't already have one
        # This prevents overwriting a correctly parsed set_syntax that includes arguments
        if not set_syntax and not cmd_key.endswith("?"):
            set_syntax = cmd_key.replace("?", "")
        if not query_syntax:
            query_syntax = cmd_key if cmd_key.endswith("?") else cmd_key + "?"
    
    # Determine commandType: "query", "set", or "both"
    if is_query_only:
        command_type = "query"
    elif set_syntax and query_syntax:
        command_type = "both"
    elif query_syntax and not set_syntax:
        command_type = "query"
    else:
        command_type = "set"
    
    cmd_entry = {
        "scpi": cmd_key,
        "description": desc,
        "conditions": cmd.get("conditions"),
        "group": group_name,
        "syntax": syntax_lines,
        "relatedCommands": cmd.get("relatedCommands"),
        "arguments": cmd.get("arguments"),
        "examples": cmd.get("examples"),
        "returns": cmd.get("returns"),
        "shortDescription": short_desc,
        "notes": cmd.get("notes", []),
        "name": name,
        "params": params,
        "example": cmd.get("examples", [{}])[0].get("scpi") if cmd.get("examples") else None,
        "_manualEntry": {
            "command": cmd_key,
            "header": cmd_key.split(" ")[0].split("?")[0],
            "mnemonics": cmd_key.split(" ")[0].split(":"),
            "commandType": command_type,
            "description": desc,
            "shortDescription": short_desc,
            "arguments": None,
            "examples": [{"description": e.get("description", ""), "codeExamples": {"scpi": {"code": e.get("scpi", "")}}} for e in (cmd.get("examples") or [])],
            "relatedCommands": cmd.get("relatedCommands") or [],
            "commandGroup": group_name,
            "syntax": {"set": set_syntax, "query": query_syntax} if set_syntax else {"query": query_syntax},
            "manualReference": {"section": group_name},
            "notes": cmd.get("notes", [])
        }
    }
    
    groups_dict[group_name]["commands"].append(cmd_entry)

# Output
output = {
    "version": "2.0",
    "manual": "4/5/6 Series MSO Programmer Manual",
    "groups": groups_dict,
    "metadata": {"total_commands": len(commands), "total_groups": len(groups_dict)}
}

output_path = os.path.join(parent_dir, "public", "commands", "mso_2_4_5_6_7.json")
print(f"\nSaving to {output_path}...")
with open(output_path, "w", encoding="utf-8") as f:
    json.dump(output, f, indent=2, ensure_ascii=False)

print(f"Done! {len(commands)} commands in {len(groups_dict)} groups")

