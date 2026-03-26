"""Quick script to inspect Word document structure"""
from docx import Document
import sys
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
DOCX_PATH = os.path.join(PROJECT_ROOT, "4-5-6_MSO_Programmer_077189801_RevA.docx")

doc = Document(DOCX_PATH)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}\n")

# Check first few tables to understand structure
print("=== First 3 Tables Structure ===")
for i, table in enumerate(doc.tables[:3]):
    print(f"\nTable {i+1}:")
    print(f"  Rows: {len(table.rows)}")
    if len(table.rows) > 0:
        print(f"  Columns: {len(table.columns)}")
        print(f"  First row cells: {len(table.rows[0].cells)}")
        
        # Show first few rows
        for row_idx, row in enumerate(table.rows[:5]):
            if len(row.cells) >= 2:
                cell1 = row.cells[0].text.strip()[:50]
                cell2 = row.cells[1].text.strip()[:50] if len(row.cells) > 1 else ""
                print(f"    Row {row_idx}: Cell1='{cell1}' | Cell2='{cell2}'")

# Check paragraphs for command patterns
print("\n=== Checking Paragraphs for Commands ===")
import re
CMD_PATTERN = re.compile(r'^([*][A-Za-z]{2,}\??)$|^([A-Za-z][A-Za-z0-9<>]*:[A-Za-z0-9<>:]+(?:\?)?)$')

cmd_count = 0
for para in doc.paragraphs[:1000]:  # Check first 1000 paragraphs
    text = para.text.strip()
    if text:
        words = text.split()
        if words:
            first_word = words[0]
            if CMD_PATTERN.match(first_word):
                cmd_count += 1
                if cmd_count <= 5:
                    print(f"  Found command: {first_word} | Rest: {text[:60]}")

print(f"\nFound {cmd_count} commands in first 1000 paragraphs")










