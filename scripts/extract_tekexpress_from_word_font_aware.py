"""
TekExpress SCPI Command Extraction Script - Font-Aware Version
Uses font detection to identify command headers, descriptions, syntax, and examples.
"""
import docx
import re
import json
import os
import sys
from typing import Dict, List, Optional, Tuple

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

# Font names to detect
TAHOMA_BOLD = "Tahoma"  # Headers and section headers
ARIAL_NARROW = "Arial Narrow"  # Descriptions, (Set)/(Query) markers, returns
COURIER_NEW = "Courier New"  # SCPI commands

# Section keywords
SECTION_KEYWORDS = {
    "syntax": "syntax",
    "command arguments": "arguments",
    "arguments": "arguments",
    "returns": "returns",
    "examples": "examples"
}

# Command pattern
TEKEXP_PATTERN = re.compile(r'TEKEXP:([A-Z]+)', re.IGNORECASE)

def extract_group_name(doc, filename: str) -> str:
    """Extract application suite name from document title or filename."""
    # Try filename first
    group_match = re.search(r'TekExpress_([A-Z0-9a-z]+)', filename, re.IGNORECASE)
    if group_match:
        return group_match.group(1)
    
    # Try document title (first few paragraphs)
    for para in doc.paragraphs[:50]:
        text = para.text.strip()
        match = re.search(r'TekExpress[®\s]+([A-Z0-9a-z]+)', text, re.IGNORECASE)
        if match:
            return match.group(1)
    
    return "TekExpress"

def get_font_name(para):
    """Get the primary font name from a paragraph - check run font first, then style."""
    # Check run font first (more specific)
    if para.runs:
        font = para.runs[0].font
        if font and font.name:
            return font.name
    # Fallback to style font
    if para.style and para.style.font and para.style.font.name:
        return para.style.font.name
    return None

def is_bold(para):
    """Check if paragraph is bold."""
    if para.style and para.style.font and para.style.font.bold:
        return True
    for run in para.runs:
        if run.bold:
            return True
    return False

def is_tahoma_header(para):
    """Check if paragraph is Tahoma header (titles and section headers)."""
    # Check style font (Tahoma) and style name (Heading 2 or Heading 5)
    style_font = None
    if para.style and para.style.font and para.style.font.name:
        style_font = para.style.font.name
    
    style_name = para.style.name if para.style else ""
    
    return (style_font and "Tahoma" in style_font and 
            ("Heading 2" in style_name or "Heading 5" in style_name))

def is_arial_narrow(para):
    """Check if paragraph is Arial Narrow (descriptions, markers)."""
    # Check both run and style font
    run_font = None
    if para.runs and para.runs[0].font and para.runs[0].font.name:
        run_font = para.runs[0].font.name
    
    style_font = None
    if para.style and para.style.font and para.style.font.name:
        style_font = para.style.font.name
    
    return (run_font and "Arial Narrow" in run_font) or (style_font and "Arial Narrow" in style_font)

def is_courier_new(para):
    """Check if paragraph is Courier New (SCPI commands)."""
    # Check run font (more specific)
    if para.runs:
        font = para.runs[0].font
        if font and font.name and "Courier New" in font.name:
            return True
    return False

def is_section_header(text: str) -> Optional[str]:
    """Check if text is a section header."""
    text_lower = text.lower().strip()
    for keyword, section_type in SECTION_KEYWORDS.items():
        if text_lower == keyword or text_lower.startswith(keyword + " "):
            return section_type
    return None

def is_tekexp_command(text: str) -> bool:
    """Check if text contains a TEKEXP command."""
    return bool(TEKEXP_PATTERN.search(text))

def extract_command_header(text: str) -> Optional[str]:
    """Extract command header from text, including subcommand."""
    match = TEKEXP_PATTERN.search(text)
    if match:
        cmd_part = text[match.start():]
        # Remove (Set) or (Query) markers
        cmd_part = re.sub(r'\s*\(Set\)\s*', '', cmd_part, flags=re.IGNORECASE)
        cmd_part = re.sub(r'\s*\(Query\)\s*', '', cmd_part, flags=re.IGNORECASE)
        # Extract command parts
        # Pattern: TEKEXP:SELECT DEVICE,"<DeviceName>" or TEKEXP:SELECT? DEVICE
        # Split by space, but handle quoted strings
        parts = []
        current = ""
        in_quotes = False
        for char in cmd_part:
            if char == '"':
                in_quotes = not in_quotes
                current += char
            elif char == ' ' and not in_quotes:
                if current.strip():
                    parts.append(current.strip())
                current = ""
            else:
                current += char
        if current.strip():
            parts.append(current.strip())
        
        if parts:
            header = parts[0].strip().replace('?', '')
            # Include subcommand if present (next uppercase word before comma or quote)
            if len(parts) > 1:
                next_part = parts[1].strip()
                # Extract just the word before comma or quote
                # "DEVICE," -> "DEVICE"
                # "DEVICE" -> "DEVICE"
                next_part = next_part.split(',')[0].split('"')[0].strip()
                # Check if it's a subcommand (uppercase, not a placeholder)
                if (next_part.isupper() and 
                    not next_part.startswith('<') and 
                    len(next_part) > 1 and
                    next_part not in ['<DeviceName>', '<SuiteName>', '<TestName>', '<Value>', '<Mode>', '<Device', '<Field>']):
                    header = f"{header} {next_part}"
            return header
    return None

def parse_argument_table(table) -> List[Dict]:
    """Parse argument table and return list of argument definitions."""
    if not table or len(table.rows) < 2:
        return []
    
    arguments = []
    header_row = [cell.text.strip() for cell in table.rows[0].cells]
    header_text = ' '.join(header_row).lower()
    
    # Check if it's an argument table
    if 'argument' not in header_text:
        return []
    
    # Simple argument table: Argument Name | Argument Type or Argument Name | Argument value
    if 'argument name' in header_text:
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                arg_name = cells[0].strip('<>')
                arg_value = cells[1].strip()
                
                # Determine type
                arg_type = "quoted_string"
                valid_values = {}
                
                # Check if it's an enumeration (has | separator or bullet points)
                if '|' in arg_value or '•' in arg_value:
                    # Extract enumeration values
                    values = []
                    # Split by | or bullet
                    parts = re.split(r'[|•]', arg_value)
                    for part in parts:
                        val = part.strip()
                        if val and val.lower() not in ['table continued…', 'table continued', 'continued…', 'continued']:
                            values.append(val)
                    
                    if values:
                        arg_type = "enumeration"
                        valid_values = {
                            "type": "enumeration",
                            "values": values,
                            "caseSensitive": False
                        }
                elif '<String>' in arg_value or 'String' in arg_value:
                    arg_type = "quoted_string"
                elif 'NR1' in arg_value.upper():
                    arg_type = "numeric"
                    valid_values = {"type": "numeric", "format": "NR1"}
                elif 'NR2' in arg_value.upper() or 'NR3' in arg_value.upper():
                    arg_type = "numeric"
                    valid_values = {"type": "numeric", "format": arg_value.upper()}
                else:
                    arg_type = "quoted_string"
                
                arguments.append({
                    "name": arg_name.lower().replace('<', '').replace('>', '').replace(' ', ''),
                    "type": arg_type,
                    "required": True,
                    "position": len(arguments),
                    "description": f"{arg_name}: {arg_value}" if arg_value else f"{arg_name}",
                    "validValues": valid_values if valid_values else None
                })
    
    return arguments

def extract_commands(doc, group_name: str) -> List[Dict]:
    """Extract all commands from the document using font detection."""
    commands = []
    current_cmd = None
    current_section = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # FIRST: Check if this is a new command title (Tahoma, Heading 2)
        # If so, finalize previous command first
        if is_tahoma_header(para) and "Heading 2" in (para.style.name if para.style else ""):
            text_lower = text.lower()
            if (text_lower.startswith('set or query') or 
                text_lower.startswith('query') or
                text_lower.startswith('generate') or
                text_lower.startswith('select')):
                
                # Finalize previous command if exists
                if current_cmd and current_cmd["header"]:
                    commands.append(current_cmd)
                
                # Start new command
                current_cmd = {
                    "header": None,  # Will be set when we find the syntax
                    "title": text,  # The Tahoma title
                    "description": "",
                    "syntax": {"set": None, "query": None},
                    "arguments": [],
                    "returns": None,
                    "examples": [],
                    "group": group_name,
                    "para_index": i
                }
                current_section = None
                continue
        
        # Check for section header (Tahoma, Heading 5)
        section_type = is_section_header(text)
        if section_type and is_tahoma_header(para) and "Heading 5" in (para.style.name if para.style else ""):
            if current_cmd:
                current_section = section_type
            continue
        
        # If we have a current command, process content
        if current_cmd:
            # Description: Arial Narrow, not a command, not a section header
            if (not current_section and 
                is_arial_narrow(para) and 
                not is_tekexp_command(text) and 
                not section_type and
                not current_cmd["description"]):
                current_cmd["description"] = text
                continue
            
            # Syntax section: Look for Courier New (SCPI commands)
            if current_section == "syntax" and is_courier_new(para):
                if is_tekexp_command(text):
                    # Check if it has (Set) or (Query) marker (Arial Narrow)
                    # The marker might be in the same paragraph or next
                    has_set = '(Set)' in text or (i + 1 < len(doc.paragraphs) and 
                                                   '(Set)' in doc.paragraphs[i + 1].text)
                    has_query = '(Query)' in text or (i + 1 < len(doc.paragraphs) and 
                                                       '(Query)' in doc.paragraphs[i + 1].text)
                    
                    # Extract syntax
                    clean_text = text.replace('(Set)', '').replace('(Query)', '').strip()
                    if has_set or '(Set)' in text:
                        current_cmd["syntax"]["set"] = clean_text
                        # Extract header from set syntax (has subcommand)
                        if not current_cmd["header"]:
                            header = extract_command_header(clean_text)
                            if header:
                                current_cmd["header"] = header
                    if has_query or '(Query)' in text:
                        current_cmd["syntax"]["query"] = clean_text
                        # If header not set yet, extract from query (but prefer set for subcommand)
                        if not current_cmd["header"]:
                            header = extract_command_header(clean_text)
                            if header:
                                current_cmd["header"] = header
                    continue
            
            # Returns section: Arial Narrow with | separator
            if current_section == "returns" and is_arial_narrow(para):
                if not current_cmd["returns"]:
                    current_cmd["returns"] = text
                continue
            
            # Examples section: Courier New for command, Arial Narrow for description
            if current_section == "examples":
                if is_courier_new(para) and is_tekexp_command(text):
                    # The description is usually in the same paragraph after "command"
                    # or in the next Arial Narrow paragraph
                    parts = text.split(' command ', 1)
                    if len(parts) == 2:
                        cmd_code = parts[0].strip()
                        desc = parts[1].strip()
                    else:
                        cmd_code = text
                        # Check next paragraph for description
                        desc = ""
                        if i + 1 < len(doc.paragraphs):
                            next_para = doc.paragraphs[i + 1]
                            if is_arial_narrow(next_para) and not is_tekexp_command(next_para.text):
                                desc = next_para.text.strip()
                    
                    current_cmd["examples"].append({
                        "description": desc or f"Example: {cmd_code}",
                        "codeExamples": {
                            "scpi": {
                                "code": cmd_code,
                                "library": "SCPI",
                                "description": "Raw SCPI command"
                            }
                        },
                        "result": None,
                        "resultDescription": desc
                    })
                    continue
    
    # Finalize last command
    if current_cmd and current_cmd["header"]:
        commands.append(current_cmd)
    
    # Now match tables to commands
    table_idx = 0
    for table in doc.tables:
        table_idx += 1
        if table_idx < 10:  # Skip document structure tables
            continue
        
        if len(table.rows) < 2:
            continue
        
        header_row = [cell.text.strip() for cell in table.rows[0].cells]
        header_text = ' '.join(header_row).lower()
        
        if 'argument' in header_text:
            args = parse_argument_table(table)
            if args:
                # Find the command that needs arguments (one without arguments yet)
                for cmd in commands:
                    if not cmd.get("arguments") or len(cmd["arguments"]) == 0:
                        cmd["arguments"] = args
                        break
    
    return commands

def post_process_to_template(commands: List[Dict]) -> Dict:
    """Transform extracted commands to match JSON template structure."""
    processed_commands = []
    
    for cmd in commands:
        header = cmd["header"]
        if not header:
            continue  # Skip commands without syntax
        
        # Generate ID
        cmd_id = header.lower().replace(':', '_').replace('?', '').replace('*', 'star').replace(' ', '_')
        
        # Extract mnemonics
        mnemonics = [m for m in header.split() if m]
        # Split by colon if present
        if ':' in header:
            parts = header.split(':')
            mnemonics = [parts[0]] + parts[1].split() if len(parts) > 1 else [parts[0]]
        
        # Determine command type
        has_set = cmd["syntax"]["set"] is not None
        has_query = cmd["syntax"]["query"] is not None
        if has_set and has_query:
            command_type = "both"
        elif has_query:
            command_type = "query"
        else:
            command_type = "set"
        
        # Get full SCPI command (prefer set, fallback to query)
        scpi = cmd["syntax"]["set"] or cmd["syntax"]["query"] or header
        
        # Short description (first sentence)
        description = cmd.get("description", "") or cmd.get("title", "")
        short_desc = description.split('.')[0][:100] if description else ""
        
        # Process arguments
        arguments = []
        if cmd.get("arguments"):
            for arg in cmd["arguments"]:
                processed_arg = {
                    "name": arg["name"],
                    "type": arg["type"],
                    "required": arg.get("required", True),
                    "position": arg.get("position", len(arguments)),
                    "description": arg.get("description", ""),
                }
                if arg.get("validValues"):
                    processed_arg["validValues"] = arg["validValues"]
                arguments.append(processed_arg)
        
        # Process query response
        query_response = None
        if cmd.get("returns"):
            returns_text = cmd["returns"]
            # Check for enumeration format (with | separator)
            if '|' in returns_text:
                values = [v.strip() for v in returns_text.split('|')]
                query_response = {
                    "type": "enumeration",
                    "format": returns_text,
                    "description": returns_text,
                    "example": values[0] if values else ""
                }
            elif '{True | False}' in returns_text or '{1 | 0}' in returns_text:
                query_response = {
                    "type": "enumeration",
                    "format": "{True | False} or {1 | 0}",
                    "description": returns_text,
                    "example": "True"
                }
            elif '<String>' in returns_text or 'String' in returns_text:
                query_response = {
                    "type": "string",
                    "format": "String",
                    "description": returns_text,
                    "example": ""
                }
            else:
                query_response = {
                    "type": "string",
                    "format": returns_text,
                    "description": returns_text,
                    "example": ""
                }
        
        # Build final command object
        processed_cmd = {
            "id": cmd_id,
            "category": "tekexpress",
            "scpi": scpi,
            "header": header,
            "mnemonics": mnemonics,
            "commandType": command_type,
            "shortDescription": short_desc,
            "description": description,
            "instruments": {
                "families": [],
                "models": [],
                "exclusions": []
            },
            "commandGroup": cmd.get("group", "TekExpress"),
            "arguments": arguments if arguments else None,
            "queryResponse": query_response,
            "syntax": {
                "set": cmd["syntax"]["set"],
                "query": cmd["syntax"]["query"]
            } if (cmd["syntax"]["set"] or cmd["syntax"]["query"]) else None,
            "codeExamples": cmd.get("examples") if cmd.get("examples") else None
        }
        
        processed_commands.append(processed_cmd)
    
    # Group commands by group name
    groups_dict = {}
    for cmd in processed_commands:
        group_name = cmd.get("commandGroup", "TekExpress")
        if group_name not in groups_dict:
            groups_dict[group_name] = {
                "name": group_name,
                "description": f"TekExpress {group_name} Automated Test Solution",
                "commands": []
            }
        groups_dict[group_name]["commands"].append(cmd)
    
    return {
        "version": "1.0",
        "manual": "TekExpress Manual",
        "groups": groups_dict
    }

def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Extract TekExpress commands from Word manual")
    parser.add_argument("--input", "-i", help="Input Word document path")
    parser.add_argument("--output", "-o", help="Output JSON path")
    parser.add_argument("--group", "-g", help="Application suite name (overrides auto-detection)")
    
    args = parser.parse_args()
    
    # Determine input file
    if args.input:
        input_file = args.input
    else:
        possible_files = [
            os.path.join(PROJECT_ROOT, "TekExpress_USB4Tx_UserManual_EN-US_077-1702-04_077170204.docx"),
        ]
        input_file = None
        for f in possible_files:
            if os.path.exists(f):
                input_file = f
                break
        
        if not input_file:
            print("ERROR: TekExpress Word document not found.")
            print("Please specify with --input or place in project root.")
            sys.exit(1)
    
    # Determine output file
    if args.output:
        output_file = args.output
    else:
        output_file = os.path.join(PROJECT_ROOT, "public", "commands", "tekexpress.json")
    
    print(f"Loading: {os.path.basename(input_file)}")
    doc = docx.Document(input_file)
    
    # Extract group name
    if args.group:
        group_name = args.group
    else:
        group_name = extract_group_name(doc, os.path.basename(input_file))
    
    print(f"Group name: {group_name}")
    
    # Extract commands
    print("Extracting commands...")
    raw_commands = extract_commands(doc, group_name)
    print(f"Found {len(raw_commands)} commands")
    
    # Debug: show what we found
    for cmd in raw_commands[:3]:
        print(f"  Command: {cmd.get('header')} - {cmd.get('title')}")
        print(f"    Description: {cmd.get('description')[:50]}...")
        print(f"    Set: {cmd['syntax']['set']}")
        print(f"    Query: {cmd['syntax']['query']}")
    
    # Post-process to template format
    print("Post-processing to template format...")
    output_data = post_process_to_template(raw_commands)
    
    # Write output
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    print(f"\nSUCCESS! Output saved to: {output_file}")
    print(f"  Commands: {len(raw_commands)}")
    print(f"  Groups: {len(output_data['groups'])}")
    
    # Print summary by group
    for group_name, group_data in output_data['groups'].items():
        print(f"    {group_name}: {len(group_data['commands'])} commands")

if __name__ == "__main__":
    main()

