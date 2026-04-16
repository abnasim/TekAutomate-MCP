"""
TekExpress SCPI Command Extraction Script
Extracts commands from TekExpress Word manuals and outputs JSON matching the template structure.
"""
import docx
import re
import json
import os
import sys
from typing import Dict, List, Optional, Tuple

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

# Section keywords
SECTION_KEYWORDS = {
    "syntax": "syntax",
    "command arguments": "arguments",
    "arguments": "arguments",
    "returns": "returns",
    "examples": "examples"
}

# Command pattern
TEKEXP_PATTERN = re.compile(r'TEKEXP:([A-Z]+)', re.IGNORECASE)

def extract_group_name(doc, filename: str) -> str:
    """Extract application suite name from document title or filename."""
    # Try filename first
    group_match = re.search(r'TekExpress_([A-Z0-9a-z]+)', filename, re.IGNORECASE)
    if group_match:
        return group_match.group(1)
    
    # Try document title (first few paragraphs)
    for para in doc.paragraphs[:50]:
        text = para.text.strip()
        # Look for "TekExpress® <SuiteName>" pattern
        match = re.search(r'TekExpress[®\s]+([A-Z0-9a-z]+)', text, re.IGNORECASE)
        if match:
            return match.group(1)
    
    return "TekExpress"

def is_section_header(text: str) -> Optional[str]:
    """Check if text is a section header."""
    text_lower = text.lower().strip()
    for keyword, section_type in SECTION_KEYWORDS.items():
        if text_lower == keyword or text_lower.startswith(keyword + " "):
            return section_type
    return None

def is_tekexp_command(text: str) -> bool:
    """Check if text contains a TEKEXP command."""
    return bool(TEKEXP_PATTERN.search(text))

def extract_command_header(text: str) -> Optional[str]:
    """Extract command header from text (e.g., 'TEKEXP:SELECT' from 'TEKEXP:SELECT DEVICE,"<DeviceName>"')."""
    match = TEKEXP_PATTERN.search(text)
    if match:
        # Get the full command up to first space or end
        cmd_part = text[match.start():]
        # Remove (Set) or (Query) markers
        cmd_part = re.sub(r'\s*\(Set\)\s*', '', cmd_part, flags=re.IGNORECASE)
        cmd_part = re.sub(r'\s*\(Query\)\s*', '', cmd_part, flags=re.IGNORECASE)
        # Extract up to first space or quote
        parts = cmd_part.split()
        if parts:
            return parts[0].strip()
    return None

def parse_argument_table(table) -> List[Dict]:
    """Parse argument table and return list of argument definitions."""
    if not table or len(table.rows) < 2:
        return []
    
    arguments = []
    header_row = [cell.text.strip() for cell in table.rows[0].cells]
    
    # Check if it's a simple argument table (Argument Name | Argument Type)
    if len(header_row) >= 2 and 'argument' in ' '.join(header_row).lower():
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                arg_name = cells[0].strip('<>')
                arg_type_str = cells[1].strip('<>')
                
                # Determine type
                arg_type = "quoted_string"
                valid_values = {}
                
                if arg_type_str.upper() in ['STRING', 'STR']:
                    arg_type = "quoted_string"
                elif 'NR1' in arg_type_str.upper():
                    arg_type = "numeric"
                    valid_values = {"type": "numeric", "format": "NR1"}
                elif 'NR2' in arg_type_str.upper() or 'NR3' in arg_type_str.upper():
                    arg_type = "numeric"
                    valid_values = {"type": "numeric", "format": arg_type_str.upper()}
                else:
                    arg_type = "quoted_string"
                
                arguments.append({
                    "name": arg_name.lower().replace('<', '').replace('>', ''),
                    "type": arg_type,
                    "required": True,
                    "position": len(arguments),
                    "description": f"{arg_name} of type {arg_type_str}",
                    "validValues": valid_values if valid_values else None
                })
    
    # Check if it's a complex enumeration table (e.g., TestName | Value)
    elif len(header_row) >= 2 and any('testname' in h.lower() or 'name' in h.lower() for h in header_row):
        # Extract all values from first column
        values = []
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if cells and cells[0]:
                # Clean up the value (remove bullet points, etc.)
                value = cells[0].strip('•').strip()
                if value and value not in values:
                    value = value.strip()
                    # Skip table continuation markers and empty values
                    if value and value.lower() not in ['table continued…', 'table continued', 'continued…', 'continued']:
                        if value not in values:
                            values.append(value)
        
        if values:
            arguments.append({
                "name": "testname" if 'testname' in header_row[0].lower() else "value",
                "type": "enumeration",
                "required": True,
                "position": 0,
                "description": f"One of: {', '.join(values[:5])}..." if len(values) > 5 else f"One of: {', '.join(values)}",
                "validValues": {
                    "type": "enumeration",
                    "values": values,
                    "caseSensitive": False
                }
            })
            
            # Check if there's a Value column (Included/Excluded)
            if len(header_row) >= 2 and 'value' in header_row[1].lower():
                # Extract unique values from second column
                value_options = set()
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2:
                        val = cells[1].strip('•').strip()
                        val = val.strip()
                    # Clean up values (remove newlines, bullets)
                    val = val.replace('\n', ' ').replace('•', '').strip()
                    if val and val.lower() not in ['table continued…', 'table continued', 'continued…', 'continued']:
                        value_options.add(val)
                
                if value_options:
                    arguments.append({
                        "name": "value",
                        "type": "enumeration",
                        "required": True,
                        "position": 1,
                        "description": f"Value: {', '.join(sorted(value_options))}",
                        "validValues": {
                            "type": "enumeration",
                            "values": sorted(list(value_options)),
                            "caseSensitive": False
                        }
                    })
    
    return arguments

def extract_commands(doc, group_name: str) -> List[Dict]:
    """Extract all commands from the document."""
    commands = []
    current_cmd = None
    current_section = None
    description_buffer = []
    pending_description = []  # Description before command header
    
    # First pass: extract commands and their sections
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Check for section header
        section_type = is_section_header(text)
        if section_type:
            if current_cmd:
                # Save previous section content
                if current_section == "description" and description_buffer:
                    current_cmd["description"] = " ".join(description_buffer)
                    description_buffer = []
                current_section = section_type
            else:
                current_section = section_type
            continue
        
        # Check if this looks like a command description (before Syntax section)
        # Pattern: "Set or query..." or "This command..."
        if not current_cmd and not is_tekexp_command(text) and not section_type:
            # Check if it's likely a description
            text_lower = text.lower()
            if (text_lower.startswith('set or query') or 
                text_lower.startswith('this command') or
                (len(text) > 20 and len(text) < 200 and not text.startswith('TEKEXP'))):
                pending_description.append(text)
            continue
        
        # Check if this is a command header (TEKEXP:... with Set/Query)
        if is_tekexp_command(text) and ('(Set)' in text or '(Query)' in text):
            # If we have a pending command, save it first
            if current_cmd:
                if description_buffer:
                    current_cmd["description"] = " ".join(description_buffer)
                    description_buffer = []
                commands.append(current_cmd)
            
            # Start new command
            header = extract_command_header(text)
            if header:
                # Use pending description if available
                desc = " ".join(pending_description) if pending_description else ""
                pending_description = []
                
                current_cmd = {
                    "header": header.replace('?', ''),  # Remove ? from header for consistency
                    "description": desc,
                    "syntax": {"set": None, "query": None},
                    "arguments": [],
                    "returns": None,
                    "examples": [],
                    "group": group_name,
                    "para_index": i  # Track paragraph index for table matching
                }
                current_section = "syntax"  # We're in syntax section now
                description_buffer = []
                
                # Extract syntax
                if '(Set)' in text:
                    current_cmd["syntax"]["set"] = text.replace('(Set)', '').strip()
                if '(Query)' in text:
                    current_cmd["syntax"]["query"] = text.replace('(Query)', '').strip()
            continue
        
        # Process content based on current section
        if current_cmd:
            if current_section == "description":
                description_buffer.append(text)
            
            elif current_section == "syntax":
                # Additional syntax lines (Set and Query on separate lines)
                if is_tekexp_command(text):
                    if '(Set)' in text:
                        current_cmd["syntax"]["set"] = text.replace('(Set)', '').strip()
                    elif '(Query)' in text:
                        current_cmd["syntax"]["query"] = text.replace('(Query)', '').strip()
            
            elif current_section == "arguments":
                # Arguments are in tables, we'll handle them separately
                pass
            
            elif current_section == "returns":
                if text and not current_cmd["returns"]:
                    current_cmd["returns"] = text
            
            elif current_section == "examples":
                if is_tekexp_command(text):
                    # Extract example command and description
                    # Format: "TEKEXP:SELECT DEVICE,"<DeviceName>" command sets..."
                    parts = text.split(' command ', 1)
                    if len(parts) == 2:
                        cmd_code = parts[0].strip()
                        desc = parts[1].strip()
                    else:
                        cmd_code = text
                        desc = ""
                    
                    current_cmd["examples"].append({
                        "description": desc or f"Example: {cmd_code}",
                        "codeExamples": {
                            "scpi": {
                                "code": cmd_code,
                                "library": "SCPI",
                                "description": "Raw SCPI command"
                            }
                        },
                        "result": None,
                        "resultDescription": desc
                    })
    
    # Save last command
    if current_cmd:
        if description_buffer:
            current_cmd["description"] = " ".join(description_buffer)
        commands.append(current_cmd)
    
    # Now process tables to find argument tables
    # Match tables to commands by finding "Command arguments" section and matching to nearest command
    # We'll track which command is expecting arguments
    cmd_expecting_args = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # If we see "Command arguments", the next table belongs to the most recent command
        if 'command arguments' in text.lower() or (text.lower() == 'arguments'):
            # Find the most recent command before this paragraph
            for cmd in reversed(commands):
                if cmd.get("para_index", 0) < i:
                    cmd_expecting_args = cmd
                    break
    
    # Now process tables
    table_idx = 0
    for table in doc.tables:
        table_idx += 1
        # Skip early tables (document structure)
        if table_idx < 10:
            continue
        
        # Check if this looks like an argument table
        if len(table.rows) < 2:
            continue
        
        header_row = [cell.text.strip() for cell in table.rows[0].cells]
        header_text = ' '.join(header_row).lower()
        
        if 'argument' in header_text or ('testname' in header_text and 'value' in header_text):
            # Parse the table
            args = parse_argument_table(table)
            if args:
                # Try to match to a command that needs arguments
                if cmd_expecting_args and not cmd_expecting_args.get("arguments"):
                    cmd_expecting_args["arguments"] = args
                else:
                    # Find command without arguments
                    for cmd in commands:
                        if not cmd.get("arguments") or len(cmd["arguments"]) == 0:
                            cmd["arguments"] = args
                            break
    
    return commands

def post_process_to_template(commands: List[Dict]) -> Dict:
    """Transform extracted commands to match JSON template structure."""
    processed_commands = []
    
    for cmd in commands:
        header = cmd["header"]
        
        # Generate ID
        cmd_id = header.lower().replace(':', '_').replace('?', '').replace('*', 'star')
        
        # Extract mnemonics
        mnemonics = [m for m in header.split(':') if m]
        
        # Determine command type
        has_set = cmd["syntax"]["set"] is not None
        has_query = cmd["syntax"]["query"] is not None
        if has_set and has_query:
            command_type = "both"
        elif has_query:
            command_type = "query"
        else:
            command_type = "set"
        
        # Get full SCPI command (prefer set, fallback to query)
        scpi = cmd["syntax"]["set"] or cmd["syntax"]["query"] or header
        
        # Short description (first sentence)
        description = cmd.get("description", "")
        short_desc = description.split('.')[0][:100] if description else ""
        
        # Process arguments
        arguments = []
        if cmd.get("arguments"):
            for arg in cmd["arguments"]:
                processed_arg = {
                    "name": arg["name"],
                    "type": arg["type"],
                    "required": arg.get("required", True),
                    "position": arg.get("position", len(arguments)),
                    "description": arg.get("description", ""),
                }
                if arg.get("validValues"):
                    processed_arg["validValues"] = arg["validValues"]
                arguments.append(processed_arg)
        
        # Process query response
        query_response = None
        if cmd.get("returns"):
            returns_text = cmd["returns"]
            # Check for boolean format
            if '{True | False}' in returns_text or '{1 | 0}' in returns_text:
                query_response = {
                    "type": "enumeration",
                    "format": "{True | False} or {1 | 0}",
                    "description": returns_text,
                    "example": "True"
                }
            elif '<String>' in returns_text:
                query_response = {
                    "type": "string",
                    "format": "String",
                    "description": returns_text,
                    "example": ""
                }
            else:
                query_response = {
                    "type": "string",
                    "format": returns_text,
                    "description": returns_text,
                    "example": ""
                }
        
        # Build final command object
        processed_cmd = {
            "id": cmd_id,
            "category": "tekexpress",
            "scpi": scpi,
            "header": header,
            "mnemonics": mnemonics,
            "commandType": command_type,
            "shortDescription": short_desc,
            "description": description,
            "instruments": {
                "families": [],
                "models": [],
                "exclusions": []
            },
            "commandGroup": cmd.get("group", "TekExpress"),
            "arguments": arguments if arguments else None,
            "queryResponse": query_response,
            "syntax": {
                "set": cmd["syntax"]["set"],
                "query": cmd["syntax"]["query"]
            } if (cmd["syntax"]["set"] or cmd["syntax"]["query"]) else None,
            "codeExamples": cmd.get("examples") if cmd.get("examples") else None
        }
        
        processed_commands.append(processed_cmd)
    
    # Group commands by group name
    groups_dict = {}
    for cmd in processed_commands:
        group_name = cmd.get("commandGroup", "TekExpress")
        if group_name not in groups_dict:
            groups_dict[group_name] = {
                "name": group_name,
                "description": f"TekExpress {group_name} Automated Test Solution",
                "commands": []
            }
        groups_dict[group_name]["commands"].append(cmd)
    
    return {
        "version": "1.0",
        "manual": "TekExpress Manual",
        "groups": groups_dict
    }

def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Extract TekExpress commands from Word manual")
    parser.add_argument("--input", "-i", help="Input Word document path")
    parser.add_argument("--output", "-o", help="Output JSON path")
    parser.add_argument("--group", "-g", help="Application suite name (overrides auto-detection)")
    
    args = parser.parse_args()
    
    # Determine input file
    if args.input:
        input_file = args.input
    else:
        # Look for TekExpress manual in project root
        possible_files = [
            os.path.join(PROJECT_ROOT, "TekExpress_USB4Tx_UserManual_EN-US_077-1702-04_077170204.docx"),
        ]
        input_file = None
        for f in possible_files:
            if os.path.exists(f):
                input_file = f
                break
        
        if not input_file:
            print("ERROR: TekExpress Word document not found.")
            print("Please specify with --input or place in project root.")
            sys.exit(1)
    
    # Determine output file
    if args.output:
        output_file = args.output
    else:
        output_file = os.path.join(PROJECT_ROOT, "public", "commands", "tekexpress.json")
    
    print(f"Loading: {os.path.basename(input_file)}")
    doc = docx.Document(input_file)
    
    # Extract group name
    if args.group:
        group_name = args.group
    else:
        group_name = extract_group_name(doc, os.path.basename(input_file))
    
    print(f"Group name: {group_name}")
    
    # Extract commands
    print("Extracting commands...")
    raw_commands = extract_commands(doc, group_name)
    print(f"Found {len(raw_commands)} commands")
    
    # Post-process to template format
    print("Post-processing to template format...")
    output_data = post_process_to_template(raw_commands)
    
    # Write output
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    print(f"\nSUCCESS! Output saved to: {output_file}")
    print(f"  Commands: {len(raw_commands)}")
    print(f"  Groups: {len(output_data['groups'])}")
    
    # Print summary by group
    for group_name, group_data in output_data['groups'].items():
        print(f"    {group_name}: {len(group_data['commands'])} commands")

if __name__ == "__main__":
    main()

